# importing Libraries

import pandas as pd
import numpy as np
import dash
import os
import re
import docx
import base64
import subprocess
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
import shutil
import socket
from dash import dcc, html, Input, Output, State, Dash
import dash_bootstrap_components as dbc
from dash_bootstrap_components._components.Container import Container
import dash_daq as daq
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime
from datetime import timedelta
from scipy.stats import percentileofscore
import datetime as dt
import calendar
from pandas.tseries.offsets import MonthEnd
from dateutil.relativedelta import relativedelta
from dash_iconify import DashIconify

pd.set_option('display.max_columns', None)


# LOAD INPUTS ##################
load_start_date = "2023-09-30"
load_end_date = "2024-1-31"

measures = [
    'MarketCap',
    'PE_Ratio',
    'EarningsYield',
    'PriceBook',
    'ReturnonTotalEquity(%)',
    'GrowthofNetIncome(%)',
    'GrowthofNetSales(%)',
    'GrowthofFreeCashFlow(%)',
    'NetProfitMargin(%)',
    'PayoutRatio',
    'TotalDebt/TotalAssets',
    'InterestCover(EBIT)',
    'ShortSell%'
]

measures_category = [
    'Size',
    'Value',
    'Value',
    'Value',
    'Value',
    'Growth',
    'Growth',
    'Growth',
    'Quality',
    'Quality',
    'Quality',
    'Quality',
    'Volatility'
]


class Portfolio:
    def __init__(self, portfolioCode):
        self.df_L1_w = pd.read_parquet('./ServerData/'+portfolioCode+'/df_L1_w.parquet')
        self.df_L1_w.index = pd.to_datetime(self.df_L1_w.index, format= '%Y-%m-%d')
        self.df_L2_w = pd.read_parquet('./ServerData/'+portfolioCode+'/df_L2_w.parquet')
        self.df_L2_w.index = pd.to_datetime(self.df_L2_w.index, format= '%Y-%m-%d')
        self.df_L3_w = pd.read_parquet('./ServerData/'+portfolioCode+'/df_L3_w.parquet')
        self.df_L3_w.index = pd.to_datetime(self.df_L3_w.index, format='%Y-%m-%d')
        self.df_L3_limits = pd.read_parquet('./ServerData/'+portfolioCode+'/df_L3_limits.parquet')
        self.df_L1_r = pd.read_parquet('./ServerData/'+portfolioCode+'/df_L1_r.parquet')
        self.df_L1_r.index = pd.to_datetime(self.df_L1_r.index, format='%Y-%m-%d')
        self.df_L1_contrib = pd.read_parquet('./ServerData/'+portfolioCode+'/df_L1_contrib.parquet')
        self.df_L1_contrib.index = pd.to_datetime(self.df_L1_contrib.index, format='%Y-%m-%d')
        self.df_L2_r = pd.read_parquet('./ServerData/'+portfolioCode+'/df_L2_r.parquet')
        self.df_L2_r.index = pd.to_datetime(self.df_L2_r.index, format= '%Y-%m-%d')
        self.df_L2_contrib = pd.read_parquet('./ServerData/'+portfolioCode+'/df_L2_contrib.parquet')
        self.df_L2_contrib.index = pd.to_datetime(self.df_L2_contrib.index, format='%Y-%m-%d')
        self.df_L3_r = pd.read_parquet('./ServerData/'+portfolioCode+'/df_L3_r.parquet')
        self.df_L3_r.index = pd.to_datetime(self.df_L3_r.index, format='%Y-%m-%d')
        self.df_L3_contrib = pd.read_parquet('./ServerData/'+portfolioCode+'/df_L3_contrib.parquet')
        self.df_L3_contrib.index = pd.to_datetime(self.df_L3_contrib.index, format='%Y-%m-%d')
        self.df_L2vsL1_relw = pd.read_parquet('./ServerData/'+portfolioCode+'/df_L2vsL1_relw.parquet')
        self.df_L2vsL1_relw.index = pd.to_datetime(self.df_L2vsL1_relw.index, format='%Y-%m-%d')
        self.df_L3vsL2_relw = pd.read_parquet('./ServerData/'+portfolioCode+'/df_L3vsL2_relw.parquet')
        self.df_L3vsL2_relw.index = pd.to_datetime(self.df_L3vsL2_relw.index, format='%Y-%m-%d')
        self.df_L3_2FAttrib = pd.read_parquet('./ServerData/'+portfolioCode+'/df_L3_2FAttrib.parquet')
        self.df_L3_2FAttrib.index = pd.to_datetime(self.df_L3_2FAttrib.index, format='%Y-%m-%d')
        self.df_L3_1FAttrib = pd.read_parquet('./ServerData/'+portfolioCode+'/df_L3_1FAttrib.parquet')
        self.df_L3_1FAttrib.index = pd.to_datetime(self.df_L3_1FAttrib.index, format='%Y-%m-%d')
        self.t_dates = pd.read_parquet('./ServerData/'+portfolioCode+'/t_dates.parquet')
        self.tME_dates = pd.read_parquet('./ServerData/'+portfolioCode+'/tME_dates.parquet')
        self.tQE_dates = pd.read_parquet('./ServerData/'+portfolioCode+'/tQE_dates.parquet')
        self.r_dates = pd.read_parquet('./ServerData/' + portfolioCode + '/r_dates.parquet')
        self.rME_dates = pd.read_parquet('./ServerData/'+portfolioCode+'/rME_dates.parquet')
        self.rQE_dates = pd.read_parquet('./ServerData/'+portfolioCode+'/rQE_dates.parquet')
        self.df_productList = pd.read_parquet('./ServerData/'+portfolioCode+'/df_productList.parquet')
        self.df_accountList = pd.read_parquet('./ServerData/' + portfolioCode + '/df_accountList.parquet')
        self.df_BM_G1 = pd.read_parquet('./ServerData/'+portfolioCode+'/df_BM_G1.parquet')
        self.df_BM_G2 = pd.read_parquet('./ServerData/'+portfolioCode+'/df_BM_G2.parquet')
        self.df_BM_G3 = pd.read_parquet('./ServerData/'+portfolioCode+'/df_BM_G3.parquet')
        self.summaryVariables = pd.read_parquet('./ServerData/'+portfolioCode+'/summaryVariables.parquet')
        self.df_Eco_USInterestRates = pd.read_parquet('./ServerData/'+portfolioCode+'/df_Eco_USInterestRates.parquet')
        # Recreate Category Group Labels for Charts
        self.portfolioCode = self.summaryVariables['portfolioCode'].iloc[0]
        self.portfolioName = self.summaryVariables['portfolioName'].iloc[0]
        self.portfolioType = self.summaryVariables['portfolioType'].iloc[0]
        self.t_StartDate = self.summaryVariables['t_StartDate'].iloc[0]
        self.t_EndDate = self.summaryVariables['t_EndDate'].iloc[0]
        self.r_StartDate = self.summaryVariables['r_StartDate'].iloc[0]
        self.bt_StartDate = self.summaryVariables['bt_StartDate'].iloc[0]
        self.groupName = self.df_BM_G1.columns[0]
        self.groupList = self.df_BM_G1[self.df_BM_G1.columns[0]].unique()



def f_get_subfolder_names(path):
    subfolder_names = []

    if os.path.exists(path) and os.path.isdir(path):
        for item in os.listdir(path):
            item_path = os.path.join(path, item)
            if os.path.isdir(item_path):
                subfolder_names.append(item)

    return subfolder_names

# IMPORT Datafiles stored on Kev's GitHUB Registry
# Must be linked to "./ServerData/"
availablePortfolios = f_get_subfolder_names('./ServerData/')

df_marketCommentary = pd.read_csv('./ServerData/MarketCommentary.csv', encoding='latin1') #utf-8-sig

# Create Portfolio class objects (import all data)
All_Portfolios = []
All_PortfolioNames = []
n = 0
for code in availablePortfolios:
    print(code)
    All_Portfolios.append(Portfolio(code))
    All_PortfolioNames.append(All_Portfolios[n].portfolioName)
    n += 1

# Initialise Globals changed by reach app calls (at bottom)
Selected_Portfolio = All_Portfolios[2]
Selected_Code = Selected_Portfolio.portfolioCode
Selected_Name = Selected_Portfolio.portfolioName
Selected_Type = Selected_Portfolio.portfolioType

Alt1_Portfolio = All_Portfolios[1]
Alt1_Code = Alt1_Portfolio.portfolioCode
Alt1_Name = Alt1_Portfolio.portfolioName
Alt1_Type = Alt1_Portfolio.portfolioType

Alt2_Portfolio = All_Portfolios[4]
Alt2_Code = Alt2_Portfolio.portfolioCode
Alt2_Name = Alt2_Portfolio.portfolioName
Alt2_Type = Alt2_Portfolio.portfolioType

text_Start_Date = load_start_date
text_End_Date = load_end_date

Alt1_Switch_On = False
Alt2_Switch_On = False

Product_List = Selected_Portfolio.df_productList.index.tolist()

dt_start_date = pd.to_datetime(text_Start_Date)
dt_end_date = pd.to_datetime(text_End_Date)
groupName = Selected_Portfolio.groupName
groupList = Selected_Portfolio.groupList

color_ACdarkblue = "#3D555E"  #BG Grey/Green
color_ACdarkblue60 = "#86959B"  #BG Grey/Green
color_ACdarkblue130 = "#223137"  #BG Grey/Green
color_ACdarkblue30 = "#C1C9CC"  #BG Grey/Green
color_ACwhite = "#E7EAEB"  #Off White
color_ACgreen = "#93F205"  #Green
color_ACgreen60 = "#C0F992"  #Green
color_ACgreen130 = "#599602"  #Green
color_ACblue = "#1DC8F2"  #Blue
color_ACblue60 = "#93DFF8"  #Blue
color_ACblue130 = "#0E7B96"  #Blue
color_ACorange = "#F27D11"  #Orange
color_ACorange60 = "#FCB384"  #Orange
color_ACorange130 = "#964B06"  #Orange

# START APP %%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%

app = dash.Dash(__name__, external_stylesheets=[dbc.themes.BOOTSTRAP, dbc.themes.MATERIA, dbc.icons.FONT_AWESOME], suppress_callback_exceptions=True)
server = app.server
port_number = 8050

def get_local_ip(port_number):
    try:
        host_name = socket.gethostname()
        local_ip = socket.gethostbyname(host_name)
        return f'http://{local_ip}:{port_number}'
    except socket.error:
        return "Couldn't get local IP address"
dashLocation = get_local_ip(port_number)
print("**** Atchison Analytics Dash App Can Be Accessed Via Local Server Running On Kev's PC Here: ")
print(dashLocation)


# %%%%%%%%%%% CORE FUNCTIONS - Calculation Return and Volatility Results

# Calculation Performance Index
def f_CalcReturnValues(df_Input, startDate, endDate):
    # example use:  returnOutput = f_CalcReturnValues(df_L3_r, dates1.loc[1,'Date'], dates1.loc[0,'Date'])
    returnOutput = 0.0
    days = 0.0
    days = ((endDate - startDate).days)

    pd.set_option('display.max_rows', None)

    if days > 0:
        returnOutput = np.prod((df_Input.loc[startDate + relativedelta(days=1):endDate] + 1)) -1 #((df_Input.loc[startDate + relativedelta(days=1):endDate] + 1).cumprod() - 1).iloc[-1]
        if startDate == endDate-timedelta(days=365): print(df_Input.loc[startDate + relativedelta(days=355):endDate])

    elif days == 0:
        returnOutput = df_Input.iloc[0]
    else:
        returnOutput = 0  # throw error here

    if days > 365: returnOutput = (((1 + returnOutput) ** (1 / (days / 365))) - 1)

    return returnOutput


def f_CalcReturnTable(df_Input, dateList):
    # example use:  f_CalcReturnTable(df_L3_r.loc[:,['IOZ','IVV']], tME_dates)
    df_Output = pd.DataFrame()

    for n in range(len(dateList)):
        if n > 0: df_Output[dateList.loc[n, 'Name']] = f_CalcReturnValues(df_Input, dateList.loc[n, 'Date'],
                                                                          dateList.loc[0, 'Date'])
    return df_Output


def f_CalcDrawdown(df_Input):
    # example use:  returnOutput = f_CalcReturnValues(df_L3_r.loc[:,['Data 1', 'Data 2']])
    # Calculate cumulative returns for each asset
    df_cumulative = (1 + df_Input).cumprod()
    # Calculate rolling maximum cumulative returns
    rolling_max = df_cumulative.expanding().max()
    # Calculate drawdowns as percentage decline from peaks
    drawdowns = (df_cumulative - rolling_max) / rolling_max * 100

    # Create the drawdown chart DataFrame
    drawdown_chart_data = pd.DataFrame()
    for col in df_Input.columns[:]:
        drawdown_chart_data[col + '_Drawdown'] = drawdowns[col]
    return drawdown_chart_data

def f_CalcRollingDrawdown(df_Input, window):
    # Calculate percentage returns for each asset
    percentage_returns = df_Input.replace([np.inf, -np.inf], np.nan).dropna()
    # Resample to monthly returns
    monthly_returns = (1 + percentage_returns).resample('M').prod() - 1

    # Initialize an empty DataFrame to store the results
    rolling_max_drawdowns = pd.DataFrame()

    # Iterate over each window period
    for start_idx in range(len(monthly_returns) - window + 1):
        end_idx = start_idx + window

        # Extract the window period returns
        window_returns = monthly_returns.iloc[start_idx:end_idx]

        # Calculate drawdown using the first function
        window_drawdowns = f_CalcDrawdown(window_returns)

        # Extract the maximum drawdown for each asset
        max_drawdowns = window_drawdowns.min().to_frame().T  # Convert to DataFrame

        # Set the end date of the window period as the index
        end_date = monthly_returns.index[end_idx - 1]
        max_drawdowns.index = [end_date]

        # Add the results to the DataFrame
        rolling_max_drawdowns = pd.concat([rolling_max_drawdowns, max_drawdowns])

    return rolling_max_drawdowns


def f_CalcRollingDailyVol(df_Input, window, trading_days_per_year):
    # Calculate percentage returns for each asset
    percentage_returns = df_Input.replace([np.inf, -np.inf], np.nan).dropna()
    # Calculate rolling volatility
    rolling_volatility = percentage_returns.rolling(window=window).std() * np.sqrt(trading_days_per_year)
    # Drop rows with NaN values (corresponding to the start of each series)
    rolling_volatility = rolling_volatility.dropna()
    return rolling_volatility

def f_CalcRollingMonthlyVol(df_Input, window, trading_months_per_year):
    # Calculate percentage returns for each asset
    percentage_returns = df_Input.replace([np.inf, -np.inf], np.nan).dropna()
    # Resample to monthly returns
    monthly_returns = (1 + percentage_returns).resample('M').prod() - 1
    # Calculate rolling volatility
    rolling_volatility = monthly_returns.rolling(window=window).std() * np.sqrt(trading_months_per_year)
    # Drop rows with NaN values (corresponding to the start of each series)
    rolling_volatility = rolling_volatility.dropna()
    return rolling_volatility

def f_CalcRollingMonthlyReturn(df_Input, window):
    percentage_returns = df_Input.replace([np.inf, -np.inf], np.nan).dropna()
    # Resample to monthly returns
    monthly_returns = (1 + percentage_returns).resample('M').prod() - 1
    days = window * 365 / 12

    if days > 0: returnOutput = (monthly_returns.rolling(window=window).apply(lambda x: (x + 1).prod() - 1, raw=True))
    if days > 365: returnOutput = (((1 + returnOutput) ** (1 / (days / 365))) - 1)

    # Drop rows with NaN values (corresponding to the start of each series)
    returnOutput = returnOutput.loc[dt_start_date:dt_end_date]
    # rolling_sharpe_ratio = rolling_sharpe_ratio.dropna()
    return returnOutput

def f_CalcRollingMonthlyAlpha(df_Input, window):

    percentage_returns = df_Input.replace([np.inf, -np.inf], np.nan).dropna()
    # Resample to monthly returns
    monthly_returns = (1 + percentage_returns).resample('M').prod() - 1
    days = window * 365 / 12
    if days > 0: returnOutput = (monthly_returns.rolling(window=window).apply(lambda x: (x + 1).prod() - 1, raw=True))
    if days > 365: returnOutput = (((1 + returnOutput) ** (1 / (days / 365))) - 1)

    monthly_returns.replace(0.0000, np.nan, inplace=True)
    monthly_returns[' vs SAA Benchmark'] = monthly_returns['P_TOTAL'] - monthly_returns['BM_G1_TOTAL']
    monthly_returns[' vs Peer Group'] = monthly_returns['P_TOTAL'] - monthly_returns['Peer_TOTAL']

    alphaOutput = monthly_returns[[' vs SAA Benchmark', ' vs Peer Group']].loc[dt_start_date:dt_end_date]


    return alphaOutput

def f_CalcRollingMonthlySharpe(df_Input, window, trading_months_per_year, risk_free_rate):
    # Calculate percentage returns for each asset
    percentage_returns = df_Input.replace([np.inf, -np.inf], np.nan).dropna()
    # Resample to monthly returns
    monthly_returns = (1 + percentage_returns).resample('M').prod() - 1
    # Calculate rolling volatility
    rolling_volatility = monthly_returns.rolling(window=window).std() * np.sqrt(trading_months_per_year)
    #rolling_volatility = rolling_volatility.iloc[-1]
    days = window*365/12

    if days > 0: returnOutput = (monthly_returns.rolling(window=window).apply(lambda x: (x + 1).prod() - 1, raw=True))

    if days > 365: returnOutput = (((1 + returnOutput) ** (1 / (days / 365))) - 1)

    rolling_mean_excess_returns = returnOutput - risk_free_rate/12

    if rolling_volatility.empty:
        return pd.DataFrame()

    rolling_sharpe_ratio = rolling_mean_excess_returns / rolling_volatility
    # Drop rows with NaN values (corresponding to the start of each series)
    rolling_sharpe_ratio = rolling_sharpe_ratio.loc[dt_start_date:dt_end_date]
    #rolling_sharpe_ratio = rolling_sharpe_ratio.dropna()
    return rolling_sharpe_ratio

def f_CalcRollingMonthlyTrackingError(df_Input, window, trading_months_per_year):

    percentage_returns = df_Input.replace([np.inf, -np.inf], np.nan).dropna()
    monthly_returns = (1 + percentage_returns).resample('M').prod() - 1
    monthly_returns.replace(0.0000, np.nan, inplace=True)

    monthly_returns['TE to Benchmark'] = monthly_returns['P_TOTAL'] - monthly_returns['BM_G1_TOTAL']
    monthly_returns['TE to Peers'] = monthly_returns['P_TOTAL'] - monthly_returns['Peer_TOTAL']

    rolling_volatility = monthly_returns.rolling(window=window).std() * np.sqrt(trading_months_per_year)
    rolling_volatility = rolling_volatility.dropna()
    rolling_volatility = rolling_volatility[['TE to Benchmark', 'TE to Peers']].loc[dt_start_date:dt_end_date]

    return rolling_volatility

def count_positive_values(x):
    return np.sum(x > 0)

def f_CalcRollingMonthlyBattingAverage(df_Input, window):
    # Calculate percentage returns for each asset
    percentage_returns = df_Input.replace([np.inf, -np.inf], np.nan).dropna()
    # Resample to monthly returns
    monthly_returns = (1 + percentage_returns).resample('M').prod() - 1
    monthly_returns.replace(0.0000, np.nan, inplace=True)

    monthly_returns['Alpha to Benchmark'] = monthly_returns['P_TOTAL'] - monthly_returns['BM_G1_TOTAL']
    monthly_returns['Alpha to Peers'] = monthly_returns['P_TOTAL'] - monthly_returns['Peer_TOTAL']

    monthly_returns = monthly_returns.dropna()
    if monthly_returns.empty:
        return pd.DataFrame()

    valid_window = min(window, monthly_returns.shape[0])

    rolling_batting_average = monthly_returns[['Alpha to Benchmark', 'Alpha to Peers']].rolling(window=window).apply(count_positive_values) / valid_window
    rolling_batting_average = rolling_batting_average.loc[dt_start_date:dt_end_date]

    return rolling_batting_average


def f_AssetClassContrib(df_Input, Input_G1_Name):
    columns_to_include = df_Input.columns[(df_Input != 0).any()].tolist()
    indices_with_G1 = Selected_Portfolio.df_productList[Selected_Portfolio.df_productList['G1'] == Input_G1_Name].index

    set1 = set(columns_to_include)
    set2 = set(indices_with_G1)
    common_elements = set1.intersection(set2)
    common_elements_list = list(common_elements)

    # Ensure that indices_with_G1 are valid indices in df_Input
    if len(common_elements_list) == 0:
        print("No valid indices found.")
        return None

    return common_elements_list


# Create Report Fucntions

def f_save_report(selected_report):
    print(f"Generating report for {selected_report}")

# Create Sidebar %%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%

left_sidebar = html.Div(
    [

        html.Span(html.I(className="fa-solid fa-thumbtack", style={"color": "#1DC8F2", "background-color": "#3D555E", "margin-left": "1rem"}),
                    id="pin-toggle-button", className="sidebar-header", n_clicks=0, style={'cursor': 'pointer'}),
        html.Div(
            [
                html.H2("Atchison Analytics", style={"color": "#1DC8F2"}),
            ],
            className="sidebar-header",
        ),
        dcc.Store(id='stored-portfolio-code', data={'key': Selected_Code}),
        dcc.Store(id='stored-alt1-switch', data={'key': Alt1_Switch_On}),
        dcc.Store(id='stored-alt2-switch', data={'key': Alt2_Switch_On}),
        html.Div(id='display-portfolio-code', style={"color": "#1DC8F2", "margin-left": "5rem"}, className="sidebar-subheader"),
        html.Hr(),
        html.Hr(style={'border-color': "#1DC8F2", 'width': '80%', 'margin': '0 auto'}),

        dbc.Nav(
            [
                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-gear me-2"),
                        html.Span("Portfolio Settings")],
                    href="/",
                    active="exact",
                ),
                html.Hr(style={'border-color': "#1DC8F2", 'width': '80%', 'margin': '0 auto'}),
                dbc.NavLink(
                    [
                        html.I(className="fas fa-home me-2"),
                        html.Span("Summary Dashboard"),
                    ],
                    href="/0-Summary",
                    active="exact",
                ),
                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-arrow-trend-up me-2"),
                        html.Span("Performance"),
                    ],
                    href="/1-Performance",
                    active="exact",
                ),
                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-face-surprise me-2"),
                        html.Span("Risk Analysis"),
                    ],
                    href="/2-Risk",
                    active="exact",
                ),
                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-chart-pie me-2"),
                        html.Span("Allocation / Exposure"),
                    ],
                    href="/3-Allocation",
                    id="menu-parent-allocation",
                    active="exact",
                ),

                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-industry me-2"),
                        html.Span("Equity Sleeve Detail"),
                    ],
                    href="/3A-Equity",
                    active="exact",
                ),

                dbc.NavLink(
                    [
                        html.I(className="fa-regular fa-credit-card me-2"),
                        html.Span("Debt Sleeve Detail"),
                    ],
                    href="/3B-Debt",
                    active="exact",
                ),
                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-coins me-2"),
                        html.Span("Alternate Sleeve Detail"),
                    ],
                    href="/3C-Alternate",
                    active="exact",
                ),
                html.Hr(style={'border-color': "#1DC8F2", 'width': '80%', 'margin': '0 auto'}),
                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-trophy me-2"),
                        html.Span("Brinson-Fachler Attribution"),
                    ],
                    href="/4-Attribution",
                    active="exact",
                ),
                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-scale-unbalanced me-2"),
                        html.Span("Contribution Analysis"),
                    ],
                    href="/5-Contribution",
                    active="exact",
                ),
                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-shapes me-2"),
                        html.Span("Portfolio Components"),
                    ],
                    href="/6-Component",
                    active="exact",
                ),
                html.Hr(style={'border-color': "#1DC8F2", 'width': '80%', 'margin': '0 auto'}),
                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-tree me-2"),
                        html.Span("ESG / Controversy"),
                    ],
                    href="/10-ESG",
                    active="exact",
                ),
                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-sack-dollar me-2"),
                        html.Span("Fee Analysis"),
                    ],
                    href="/11-Fees",
                    active="exact",
                ),
                html.Hr(style={'border-color': "#1DC8F2", 'width': '80%', 'margin': '0 auto'}),
                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-landmark me-2"),
                        html.Span("Market Valuation Analysis"),
                    ],
                    href="/20-Markets",
                    active="exact",
                ),
                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-file-lines me-2"),
                        html.Span("Report Generator"),
                    ],
                    href="/21-Reports",
                    active="exact",
                ),
                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-cloud-arrow-down me-2"),
                        html.Span("HTML Download"),
                    ],
                    href="/22-Download",
                    active="exact",
                ),
                html.Hr(style={'border-color': "#1DC8F2", 'width': '80%', 'margin': '0 auto'}),
                dbc.NavLink(
                    [
                        html.I(className="fa-solid fa-circle-info me-2"),
                        html.Span("Need Help?"),
                    ],
                    href="/30-Help",
                    active="exact",
                ),

            ],
            vertical=True,
            pills=True,
        ),
    ],
    className="sidebar",
    id="sidebar-left-id",
)


right_sidebar = html.Div(
    [
        html.Div(
            [
                html.H2("Settings", style={"color": "#1DC8F2"}),
            ],
            className="right_sidebar-header",
        ),
        html.I(className="fa-solid fa-gear me-2"),
        html.Hr(),
        html.Hr(style={'border-color': "#1DC8F2", 'width': '80%', 'margin': '0 auto'}),
        html.Hr(),
        html.Hr(style={'border-color': "#1DC8F2", 'width': '80%', 'margin': '0 auto'}),

    ],
    className="right_sidebar",
    id="sidebar-right-id",
)

# MAin AREA FIGURE FUNCTIONS

def f_create_SUMMARY_REPORT_HTML(in_df_marketCommentary):
    print("Create report")

def f_create_3DSURFACE_figure(df_input, in_title, in_z_title, in_y_title, in_x_title, in_height):

    custom_colors = {Selected_Code: color_ACdarkblue, 'Peer Group': color_ACblue60, 'Inflation': color_ACorange60,
                     'Objective': color_ACorange60, 'SAA Benchmark': color_ACgreen60,
                     'Growth': color_ACdarkblue, 'Defensive': color_ACdarkblue130,
                     'Australian Shares': color_ACgreen, 'International Shares': color_ACgreen130,
                     'Real Assets': color_ACorange, 'Alternatives': color_ACorange130,
                     'Long Duration': color_ACblue, 'Floating Rate': color_ACblue130,
                     'Cash': color_ACblue60}
    try:
        z = df_input.values.T
        y = pd.to_numeric(df_input.columns)
        x = df_input.index

        figure_out = go.Figure(data=[go.Surface(z=z, x=x, y=y)])

        figure_out.update_layout(
            title=in_title,
            height=in_height if in_height is not None else 800,
            margin=dict(r=0, l=0, b=0),  # Reduce right margin to maximize visible area
            images=[dict(
                source=f"data:image/png;base64,{base64.b64encode(open('./assets/atchisonlogo.png', 'rb').read()).decode()}",
                xref="paper", yref="paper",
                x=0.98, y=0.02,
                sizex=0.2, sizey=0.2,
                xanchor="right", yanchor="bottom",
                layer="below"
            )],
            legend=dict(
                orientation="h",
                yanchor="top",
                y=-0.3,
                xanchor="center",
                x=0.5,
                title=None,
                font=dict(size=11)
            ),
        )

        figure_out.update_layout(scene=dict(xaxis=dict(title_text=in_x_title), yaxis=dict(title_text=in_y_title), zaxis=dict(title_text=in_z_title)))

        return figure_out

    except Exception as e:
        print(f"An error occurred 3DSurface {in_title}: {e}")
        # Handle the error as needed
        return []  # or any default return value

def f_create_LINE_figure(df_input, in_title, in_y_title, in_x_title, in_height):
    custom_colors = {Selected_Code: color_ACdarkblue, 'Peer Group': color_ACblue60, 'Inflation': color_ACorange60,
                     'Objective': color_ACorange60, 'SAA Benchmark': color_ACgreen60,
                     'Growth': color_ACdarkblue, 'Defensive': color_ACdarkblue130,
                     'Australian Shares': color_ACgreen, 'International Shares': color_ACgreen130,
                     'Real Assets': color_ACorange, 'Alternatives': color_ACorange130,
                     'Long Duration': color_ACblue, 'Floating Rate': color_ACblue130,
                     'Cash': color_ACblue60}

    try:
        figure_out = px.line(
            df_input,
            x=df_input.index,
            y=[c for c in df_input.columns if c is not None],
            labels={"x": "Date", "y": "Values"},
            template="plotly_white",
            color_discrete_map=custom_colors
        )
        figure_out.update_layout(
            title=in_title,
            yaxis_title=in_y_title,
            xaxis_title=in_x_title,
            height=in_height if in_height is not None else 800,
            margin=dict(r=0, l=0, b=0),  # Reduce right margin to maximize visible area
            images=[dict(
                source="../assets/atchisonlogo.png",
                xref="paper", yref="paper",
                x=0.98, y=0.02,
                sizex=0.1, sizey=0.1,
                xanchor="right", yanchor="bottom",
                layer="below"
            )],
            legend=dict(
                orientation="h",
                yanchor="top",  # Change this to "top" to move the legend below the chart
                y=-0.3,  # Adjust the y value to position the legend below the chart
                xanchor="center",  # Center the legend horizontally
                x=0.5,  # Center the legend horizontally
                title=None,
                font=dict(size=11)
            ),
        )

        return figure_out
    except Exception as e:
        print(f"An error occurred in f_create_LINE_figure {in_title}: {e}")
        # Handle the error as needed
        return []  # or any default return value



def f_create_BAR_figure(df_input, in_type, in_title, in_y_title, in_x_title, in_height):
    custom_colors = {Selected_Code: color_ACdarkblue, 'Peer Group': color_ACblue60, 'Inflation': color_ACorange60,
                     'Objective': color_ACorange60, 'SAA Benchmark': color_ACgreen60,
                     'Growth': color_ACdarkblue, 'Defensive': color_ACdarkblue130,
                     'Australian Shares': color_ACgreen, 'International Shares': color_ACgreen130,
                     'Real Assets': color_ACorange, 'Alternatives': color_ACorange130,
                     'Long Duration': color_ACblue, 'Floating Rate': color_ACblue130,
                     'Cash': color_ACblue60}

    try:
        figure_out = px.bar(
            df_input,
            x=df_input.index,
            y=[c for c in df_input.columns if c is not None],
            labels={"x": "Date", "y": "Values"},
            template="plotly_white",
            barmode=in_type,
            color_discrete_map = custom_colors
        )
        figure_out.update_layout(
            title=in_title,
            yaxis_title=in_y_title,
            xaxis_title=in_x_title,
            height=in_height if in_height is not None else 800,
            margin=dict(r=0, l=0),  # Reduce right margin to maximize visible area
            images=[dict(
                source="../assets/atchisonlogo.png",
                xref="paper", yref="paper",
                x=0.98, y=1.00,
                sizex=0.1, sizey=0.1,
                xanchor="right", yanchor="top",
                layer="below"
            )],
            legend=dict(
                orientation="h",
                yanchor="top",  # Change this to "top" to move the legend below the chart
                y=-0.3,  # Adjust the y value to position the legend below the chart
                xanchor="center",  # Center the legend horizontally
                x=0.5,  # Center the legend horizontally
                title=None,
                font=dict(size=11)
            ),
        )

        return figure_out
    except Exception as e:
        print(f"An error occurred in f_create_BAR_figure {in_title}: {e}")
        # Handle the error as needed
        return []  # or any default return value


def f_create_COLORBAR_figure(df_input, in_type, in_x, in_y, in_color, in_title, in_y_title, in_x_title, in_height):
    custom_colors = {Selected_Code: color_ACdarkblue, 'Peer Group': color_ACblue60, 'Inflation': color_ACorange60,
                     'Objective': color_ACorange60, 'SAA Benchmark': color_ACgreen60,
                     'Growth': color_ACdarkblue, 'Defensive': color_ACdarkblue130,
                     'Australian Shares': color_ACgreen, 'International Shares': color_ACgreen130,
                     'Real Assets': color_ACorange, 'Alternatives': color_ACorange130,
                     'Long Duration': color_ACblue, 'Floating Rate': color_ACblue130,
                     'Cash': color_ACblue60}
    try:
        figure_out = px.bar(
            df_input,
            x=in_x,
            y=in_y,
            labels={"x": "Date", "y": "Values"},
            template="plotly_white",
            barmode=in_type,
            color=in_color
        )
        figure_out.update_layout(
            title=in_title,
            yaxis_title=in_y_title,
            xaxis_title=in_x_title,
            height=in_height if in_height is not None else 800,
            margin=dict(r=0, l=0),  # Reduce right margin to maximize visible area
            images=[dict(
                source="../assets/atchisonlogo.png",
                xref="paper", yref="paper",
                x=0.98, y=1.00,
                sizex=0.1, sizey=0.1,
                xanchor="right", yanchor="top",
                layer="below"
            )],
            legend=dict(
                orientation="h",
                yanchor="top",  # Change this to "top" to move the legend below the chart
                y=-0.3,  # Adjust the y value to position the legend below the chart
                xanchor="center",  # Center the legend horizontally
                x=0.5,  # Center the legend horizontally
                title=None,
                font=dict(size=11)
            ),
            yaxis=dict(
                range=[-50, 50],
            )
        )

        return figure_out
    except Exception as e:
        print(f"An error occurred in f_create_COLORBAR_figure {in_title}: {e}")
        # Handle the error as needed
        return []  # or any default return value

def f_create_WATERFALL_figure(df_input, in_x, in_y, in_title, in_y_title, in_x_title, in_height, y_min, y_max):
    custom_colors = {Selected_Code: color_ACdarkblue, 'Peer Group': color_ACblue60, 'Inflation': color_ACorange60,
                     'Objective': color_ACorange60, 'SAA Benchmark': color_ACgreen60,
                     'Growth': color_ACdarkblue, 'Defensive': color_ACdarkblue130,
                     'Australian Shares': color_ACgreen, 'International Shares': color_ACgreen130,
                     'Real Assets': color_ACorange, 'Alternatives': color_ACorange130,
                     'Long Duration': color_ACblue, 'Floating Rate': color_ACblue130,
                     'Cash': color_ACblue60}
    try:
        figure_out = go.Figure(go.Waterfall(
            x=df_input[in_x],
            y=df_input[in_y],
            textposition="outside",
            connector={"line": {"color": color_ACdarkblue}},
        ))
        layout_params = {
            "title": in_title,
            "xaxis_title": in_x_title,
            "height": in_height if in_height is not None else 800,
            "margin": dict(r=0, l=0),  # Reduce right margin to maximize visible area
            "images": [{
                "source": "../assets/atchisonlogo.png",
                "xref": "paper", "yref": "paper",
                "x": 0.98, "y": 0.02,
                "sizex": 0.1, "sizey": 0.1,
                "xanchor": "right", "yanchor": "bottom",
                "layer": "below"
            }],
        }
        if y_min is not None and y_max is not None:
            layout_params["yaxis"] = dict(
                range=[y_min, y_max],
                title=in_y_title,
            )
        figure_out.update_layout(**layout_params)

        return figure_out
    except Exception as e:
        print(f"An error occurred in f_create_WATERFALL_figure {in_title}: {e}")
        # Handle the error as needed
        return []  # or any default return value




def f_create_RANGE_figure(df_input, in_title, in_y_title, in_x_title, in_height):
    custom_colors = {Selected_Code: color_ACdarkblue, 'Peer Group': color_ACblue60, 'Inflation': color_ACorange60,
                     'Objective': color_ACorange60, 'SAA Benchmark': color_ACgreen60,
                     'Growth': color_ACdarkblue, 'Defensive': color_ACdarkblue130,
                     'Australian Shares': color_ACgreen, 'International Shares': color_ACgreen130,
                     'Real Assets': color_ACorange, 'Alternatives': color_ACorange130,
                     'Long Duration': color_ACblue, 'Floating Rate': color_ACblue130,
                     'Cash': color_ACblue60}
    try:
        figure_out = px.bar(
            df_input,
            x='Group Value', y=['Min', 'Max-Min'],
            template="plotly_white",
        )

        figure_out.update_traces(marker_color='#3D555E', width=0.3, opacity=0,
                                 selector=dict(name='Min'), showlegend=False)  # Set color, width, and opacity for 'Min' bars

        figure_out.update_traces(marker_color='#3D555E', width=0.3,
                                 selector=dict(name='Max-Min'), showlegend=False)  # Set color and width for 'Max-Min' bars

        scatter_fig = px.scatter(df_input, x='Group Value', y='Current',
                                 title='Current', color_discrete_sequence=['#1DC8F2'])
        for trace in scatter_fig.data:
            figure_out.add_trace(trace)

        figure_out.update_layout(
            title=in_title,
            yaxis_title=in_y_title,
            xaxis_title=in_x_title,
            height=in_height,
            margin=dict(r=0, l=0),  # Reduce right margin to maximize visible area
            images=[dict(
                source="../assets/atchisonlogo.png",
                xref="paper", yref="paper",
                x=0.98, y=1.00,
                sizex=0.1, sizey=0.1,
                xanchor="right", yanchor="top",
                layer="below"
            )],
            legend=None,
        )

        return figure_out
    except Exception as e:
        print(f"An error occurred in f_create_RANGE_figure {in_title}: {e}")
        # Handle the error as needed
        return []  # or any default return value


def f_create_SCATTER_figure(df_input, in_averages, in_x, in_y, in_size, in_color, in_title, in_y_title, in_x_title, in_height, in_dot_scale, x_range=None, y_range=None):
    custom_colors = {Selected_Code: color_ACdarkblue, 'Peer Group': color_ACblue60, 'Inflation': color_ACorange60,
                     'Objective': color_ACorange60, 'SAA Benchmark': color_ACgreen60,
                     'Growth': color_ACdarkblue, 'Defensive': color_ACdarkblue130,
                     'Australian Shares': color_ACgreen, 'International Shares': color_ACgreen130,
                     'Real Assets': color_ACorange, 'Alternatives': color_ACorange130,
                     'Long Duration': color_ACblue, 'Floating Rate': color_ACblue130,
                     'Cash': color_ACblue60}
    try:
        figure_out = px.scatter(
            df_input,
            x=in_x,
            y=in_y,
            size=in_size,
            size_max=20*in_dot_scale,
            color=in_color,
            hover_data=['Name', 'Code', 'LastPrice']+measures,
            template="plotly_white"
        )

        # Update x-axis and y-axis ranges if specified
        if x_range is not None:
            figure_out.update_xaxes(range=x_range)
        if y_range is not None:
            figure_out.update_yaxes(range=y_range)

        valid_x = re.sub(r'[^a-zA-Z0-9]', '', in_x)
        valid_y = re.sub(r'[^a-zA-Z0-9]', '', in_y)
        selected_trace = go.Scatter(x=[in_averages[f"selected_avg_{valid_x}"]],
                                          y=[in_averages[f"selected_avg_{valid_y}"]],
                                          mode='markers+text',
                                          marker=dict(size=12, color='black', symbol="cross",
                                                      line=dict(color='black', width=1)),
                                          text='Weighted Portfolio',
                                          textposition='bottom right',
                                          showlegend=False)

        valid_BMx = re.sub(r'[^a-zA-Z0-9]', '', in_x)
        valid_BMy = re.sub(r'[^a-zA-Z0-9]', '', in_y)
        BM_trace = go.Scatter(x=[in_averages[f"BM_avg_{valid_BMx}"]], y=[in_averages[f"BM_avg_{valid_BMy}"]],
                                    mode='markers+text',
                                    marker=dict(size=12, color='black', symbol="star",
                                                line=dict(color='black', width=1)),
                                    text='Weighted Benchmark',
                                    textposition='bottom right',
                                    showlegend=False)
        figure_out.add_trace(selected_trace)
        figure_out.add_trace(BM_trace)
        figure_out.update_layout(
            title=in_title,
            yaxis_title=in_y_title if in_y_title is not None else in_y,
            xaxis_title=in_x_title if in_x_title is not None else in_x,
            height=in_height if in_height is not None else 800,
            margin=dict(r=0, l=0),  # Reduce right margin to maximize visible area
            images=[dict(
                source="../assets/atchisonlogo.png",
                xref="paper", yref="paper",
                x=0.98, y=0.02,
                sizex=0.1, sizey=0.1,
                xanchor="right", yanchor="bottom",
                layer="below"
            )],
        )

        return figure_out
    except Exception as e:
        print(f"An error occurred in f_create_SCATTER_figure {in_title}: {e}")
        # Handle the error as needed
        return []  # or any default return value



def f_create_PIE_figure(df_input, in_values, in_names, in_title, in_height):
    custom_colors = {Selected_Code: color_ACdarkblue, 'Peer Group': color_ACblue60, 'Inflation': color_ACorange60,
                     'Objective': color_ACorange60, 'SAA Benchmark': color_ACgreen60,
                     'Growth': color_ACdarkblue, 'Defensive': color_ACdarkblue130,
                     'Australian Shares': color_ACgreen, 'International Shares': color_ACgreen130,
                     'Real Assets': color_ACorange, 'Alternatives': color_ACorange130,
                     'Long Duration': color_ACblue, 'Floating Rate': color_ACblue130,
                     'Cash': color_ACblue60}

    try:
        figure_out = px.pie(
            df_input,
            values=in_values,
            names=in_names,
            template="plotly_white",
            color=in_names,  # Use G1 column values for color assignment
            color_discrete_map=custom_colors
        )
        figure_out.update_layout(
            title=in_title,
            height=in_height,
            margin=dict(r=0, l=0, b=0),  # Reduce right margin to maximize visible area
            images=[dict(
                source="../assets/atchisonlogo.png",
                xref="paper", yref="paper",
                x=0.98, y=0.02,
                sizex=0.1, sizey=0.1,
                xanchor="right", yanchor="bottom",
                layer="below"
            )],
            legend=dict(
                orientation="h",
                yanchor="top",  # Change this to "top" to move the legend below the chart
                y=-0.3,  # Adjust the y value to position the legend below the chart
                xanchor="center",  # Center the legend horizontally
                x=0.5,  # Center the legend horizontally
                title=None,
                font=dict(size=11)
            ),
        )
        return figure_out
    except Exception as e:
        print(f"An error occurred in f_create_PIE_figure {in_title}: {e}")
        # Handle the error as needed
        return []  # or any default return value

def f_create_SUNBURST_figure(df_input, in_path, in_names, in_values, in_title, in_height):
    custom_colors = {Selected_Code: color_ACdarkblue, 'Peer Group': color_ACblue60, 'Inflation': color_ACorange60,
                     'Objective': color_ACorange60, 'SAA Benchmark': color_ACgreen60,
                     'Growth': color_ACdarkblue, 'Defensive': color_ACdarkblue130,
                     'Australian Shares': color_ACgreen, 'International Shares': color_ACgreen130,
                     'Real Assets': color_ACorange, 'Alternatives': color_ACorange130,
                     'Long Duration': color_ACblue, 'Floating Rate': color_ACblue130,
                     'Cash': color_ACblue60}

    try:

        figure_out = px.sunburst(
            df_input,
            path=in_path,
            names=in_names,
            values=in_values,
            template="plotly_white",
            color=in_path[0],  # Use G1 column values for color assignment
            color_discrete_map=custom_colors
        )
        figure_out.update_layout(
            title=in_title,
            height=in_height,
            margin=dict(r=0, l=0, b=0),  # Reduce right margin to maximize visible area
            images=[dict(
                source="../assets/atchisonlogo.png",
                xref="paper", yref="paper",
                x=0.98, y=0.02,
                sizex=0.1, sizey=0.1,
                xanchor="right", yanchor="bottom",
                layer="below"
            )],
        )
        return figure_out

    except Exception as e:
        print(f"An error occurred in f_create_SUNBURST_figure: {e}")
        # Handle the error as needed
        return []  # or any default return value

def f_create_POLAR_figure(df_input, in_Portfolioname, in_BMname, in_title, in_height):
    custom_colors = {Selected_Code: color_ACdarkblue, 'Peer Group': color_ACblue60, 'Inflation': color_ACorange60,
                     'Objective': color_ACorange60, 'SAA Benchmark': color_ACgreen60,
                     'Growth': color_ACdarkblue, 'Defensive': color_ACdarkblue130,
                     'Australian Shares': color_ACgreen, 'International Shares': color_ACgreen130,
                     'Real Assets': color_ACorange, 'Alternatives': color_ACorange130,
                     'Long Duration': color_ACblue, 'Floating Rate': color_ACblue130,
                     'Cash': color_ACblue60}
    try:
        # Now create Polar dataframe sets for summary chart
        figure_out = go.Figure()

        figure_out.add_trace(go.Scatterpolar(
            r=df_input['Selected MCap Normalized'],
            theta=df_input['Measure'],
            hovertext=df_input['Benchmark Avg'],
            fill='toself',
            name='MCap Weighted: ' + in_Portfolioname
        ))

        figure_out.add_trace(go.Scatterpolar(
            r=df_input['Benchmark MCap Normalized'],
            theta=df_input['Measure'],
            hovertext=df_input['Benchmark Avg'],
            fill='toself',
            name='MCap Weighted: ' + in_BMname
        ))

        figure_out.add_trace(go.Scatterpolar(
            r=df_input['Selected EW Normalized'],
            theta=df_input['Measure'],
            hovertext=df_input['Selected Avg'],
            fill='toself',
            name='Equal Weighted: ' + in_Portfolioname,
            visible='legendonly'
        ))

        figure_out.add_trace(go.Scatterpolar(
            r=df_input['Benchmark EW Normalized'],
            theta=df_input['Measure'],
            hovertext=df_input['Benchmark Avg'],
            fill='toself',
            name='Equal Weighted: ' + in_BMname,
            visible='legendonly'
        ))

        figure_out.update_layout(
            height=in_height,
            polar=dict(
                radialaxis=dict(
                    range=[-50, 50]  # Set the range to always show 0 to 100
                )
            ),
            title={
                "text": f"As at {dt_end_date:%d-%b-%Y}",
                "font": {"size": 11}  # Adjust the font size as needed
            },
            margin=dict(r=0, l=0),  # Reduce right margin to maximize visible area
            images=[dict(
                source="../assets/atchisonlogo.png",
                xref="paper", yref="paper",
                x=0.98, y=0.02,
                sizex=0.1, sizey=0.1,
                xanchor="right", yanchor="bottom",
                layer="below"
            )],
        )

        return figure_out

    except Exception as e:
        print(f"An error occurred in f_create_POLAR_figure {e}")
        # Handle the error as needed
        return []  # or any default return value


def f_FILL_1perf(Local_Portfolio):
    try:
        df_1perf_daily = pd.concat(
            [Local_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date, ['P_' + groupName + '_' + n]] * 100 for n in
             groupList], axis=1)
        df_1perf_daily.columns = groupList

        df_1perf_total = (((Local_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date,
                            ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL', 'Obj_TOTAL']] + 1).cumprod() - 1) * 100)
        df_1perf_total.columns = [Selected_Code, 'SAA Benchmark', 'Peer Group', 'Objective']

        print("Here,,,")
        print(df_1perf_total)
        if Alt1_Switch_On != False:
            a1 = (((Alt1_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date, ['P_TOTAL']] + 1).cumprod() - 1) * 100)
            a1.columns = ['Alt 1 (' + Alt1_Code + ')']
            df_1perf_total = pd.concat([df_1perf_total, a1], axis=1)
        if Alt2_Switch_On != False:
            a2 = (((Alt2_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date, ['P_TOTAL']] + 1).cumprod() - 1) * 100)
            a2.columns = ['Alt 2 (' + Alt2_Code + ')']
            df_1perf_total = pd.concat([df_1perf_total, a2], axis=1)

        df_1perf_tSet = (f_CalcReturnTable(
            Local_Portfolio.df_L3_r.loc[:, ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL', 'Obj_TOTAL']],
            Local_Portfolio.t_dates) * 100).T


        df_1perf_tSet.columns = [Selected_Code, 'SAA Benchmark', 'Peer Group', 'Objective']

        if Alt1_Switch_On != False:
            a1 = (f_CalcReturnTable(Alt1_Portfolio.df_L3_r.loc[:, ['P_TOTAL']], Local_Portfolio.t_dates) * 100).T
            a1.columns = ['Alt 1 (' + Alt1_Code + ')']
            df_1perf_tSet = pd.concat([df_1perf_tSet, a1], axis=1)
        if Alt2_Switch_On != False:
            a2 = (f_CalcReturnTable(Alt2_Portfolio.df_L3_r.loc[:, ['P_TOTAL']], Local_Portfolio.t_dates) * 100).T
            a2.columns = ['Alt 2 (' + Alt2_Code + ')']
            df_1perf_tSet = pd.concat([df_1perf_tSet, a2], axis=1)

        df_1perf_tMESet = (f_CalcReturnTable(
            Local_Portfolio.df_L3_r.loc[:, ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL', 'Obj_TOTAL']],
            Local_Portfolio.tME_dates) * 100).T
        df_1perf_tMESet.columns = [Selected_Code, 'SAA Benchmark', 'Peer Group', 'Objective']

        if Alt1_Switch_On != False:
            a1 = (f_CalcReturnTable(Alt1_Portfolio.df_L3_r.loc[:, ['P_TOTAL']], Local_Portfolio.tME_dates) * 100).T
            a1.columns = ['Alt 1 (' + Alt1_Code + ')']
            df_1perf_tMESet = pd.concat([df_1perf_tMESet, a1], axis=1)
        if Alt2_Switch_On != False:
            a2 = (f_CalcReturnTable(Alt2_Portfolio.df_L3_r.loc[:, ['P_TOTAL']], Local_Portfolio.tME_dates) * 100).T
            a2.columns = ['Alt 2 (' + Alt2_Code + ')']
            df_1perf_tMESet = pd.concat([df_1perf_tMESet, a2], axis=1)

        print(Local_Portfolio.rME_dates)

        df_1perf_rMESet = (f_CalcReturnTable(

            Local_Portfolio.df_L3_r.loc[:, ['P_TOTAL', 'Peer_TOTAL', 'Obj_TOTAL']],
            Local_Portfolio.rME_dates) * 100).T
        df_1perf_rMESet.columns = [Selected_Code, 'Peer Group', 'Inflation']

        df_1perf_rMESet["Outperformance vs Peers"] = df_1perf_rMESet[Selected_Code] - df_1perf_rMESet['Peer Group']
        df_1perf_rMESet["Outperformance vs Inflation"] = df_1perf_rMESet[Selected_Code] - df_1perf_rMESet['Inflation']

        if Alt1_Switch_On != False:
            a1 = (f_CalcReturnTable(Alt1_Portfolio.df_L3_r.loc[:, ['P_TOTAL']], Local_Portfolio.rME_dates) * 100).T
            a1.columns = ['Alt 1 (' + Alt1_Code + ')']
            df_1perf_rMESet = pd.concat([df_1perf_rMESet, a1], axis=1)
        if Alt2_Switch_On != False:
            a2 = (f_CalcReturnTable(Alt2_Portfolio.df_L3_r.loc[:, ['P_TOTAL']], Local_Portfolio.rME_dates) * 100).T
            a2.columns = ['Alt 2 (' + Alt2_Code + ')']
            df_1perf_rMESet = pd.concat([df_1perf_rMESet, a2], axis=1)


        df_1perf_tQESet = (f_CalcReturnTable(
            Local_Portfolio.df_L3_r.loc[:, ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL', 'Obj_TOTAL']],
            Local_Portfolio.tQE_dates) * 100).T

        df_1perf_tQESet.columns = [Selected_Code, 'SAA Benchmark', 'Peer Group', 'Objective']

        if Alt1_Switch_On != False:
            a1 = (f_CalcReturnTable(Alt1_Portfolio.df_L3_r.loc[:, ['P_TOTAL']], Local_Portfolio.tQE_dates) * 100).T
            a1.columns = ['Alt 1 (' + Alt1_Code + ')']
            df_1perf_tQESet = pd.concat([df_1perf_tQESet, a1], axis=1)

        if Alt2_Switch_On != False:
            a2 = (f_CalcReturnTable(Alt2_Portfolio.df_L3_r.loc[:, ['P_TOTAL']], Local_Portfolio.tQE_dates) * 100).T
            a2.columns = ['Alt 2 (' + Alt2_Code + ')']
            df_1perf_tQESet = pd.concat([df_1perf_tQESet, a2], axis=1)

        return df_1perf_daily, df_1perf_total, df_1perf_tSet, df_1perf_tMESet, df_1perf_tQESet, df_1perf_rMESet

    except Exception as e:
        print(f"An error occurred in f_create_FILL_1: {e}")
        # Handle the error as needed
        return []  # or any default return value



def f_FILL_3alloc(Local_Portfolio):
    try:
        df_3alloc_sleeves = pd.concat(
            [Local_Portfolio.df_L3_w.loc[dt_end_date:dt_end_date, ['P_' + groupName + '_' + n]].T for n in groupList],
            axis=1)
        df_3alloc_sleeves.index = groupList
        df_3alloc_sleeves['Current'] = df_3alloc_sleeves.sum(axis=1)
        df_3alloc_sleeves = df_3alloc_sleeves[['Current']]
        df_3alloc_sleeves.reset_index(inplace=True)
        df_3alloc_sleeves.columns = ['GroupValue', 'Current']

        df_3alloc_BMsleeves = pd.concat(
            [Local_Portfolio.df_L2_w.loc[dt_end_date:dt_end_date, ['BM_' + groupName + '_' + n]].T for n in groupList],
            axis=1)
        df_3alloc_BMsleeves.index = groupList
        df_3alloc_BMsleeves['Current'] = df_3alloc_BMsleeves.sum(axis=1)
        df_3alloc_BMsleeves = df_3alloc_BMsleeves[['Current']]
        df_3alloc_BMsleeves.reset_index(inplace=True)
        df_3alloc_BMsleeves.columns = ['GroupValue', 'Current']

        # Below is dependent on df_3alloc_sleeves
        row_values = []
        allrows_values = []
        group_df = Local_Portfolio.df_L3_limits[Local_Portfolio.df_L3_limits['Group'] == groupName]
        for n, element in enumerate(groupList):
            # Filter the DataFrame for the current group
            group_df2 = group_df[group_df['GroupValue'] == groupList[n]]
            row_values.append(groupList[n])
            row_values.append(df_3alloc_sleeves.loc[n, "Current"])
            # Check if there are any rows for the current group
            if not group_df2.empty:
                # Get the minimum value from the 'Min' column of the filtered DataFrame
                row_values.append(group_df2['Min'].min())
                row_values.append(group_df2['Max'].max())
            else:
                # If no rows are found for the current group, append None to the list
                row_values.append(0)
                row_values.append(100)

            allrows_values.append(row_values)
            row_values = []

        column_names = ['Group Value', 'Current', 'Min', 'Max']
        df_3alloc_sleeve_ranges = pd.DataFrame(allrows_values, columns=column_names)
        df_3alloc_sleeve_ranges['Max-Min'] = df_3alloc_sleeve_ranges['Max'] - df_3alloc_sleeve_ranges['Min']

        df_3alloc_OWUW = pd.concat([Local_Portfolio.df_L3vsL2_relw.loc[dt_start_date:dt_end_date,
                                    ['P_' + groupName + '_' + n]] for n in groupList], axis=1)
        df_3alloc_OWUW.columns = groupList

        df_3alloc_weights = pd.concat([Local_Portfolio.df_L3_w.loc[dt_start_date:dt_end_date,
                                       ['P_' + groupName + '_' + n]] for n in groupList], axis=1)
        df_3alloc_weights.columns = groupList

        df_3alloc_mgr_level = Local_Portfolio.df_L3_w.loc[dt_end_date:dt_end_date,
                              Local_Portfolio.df_L3_w.columns.isin(Product_List)].tail(1)
        df_3alloc_mgr_level = df_3alloc_mgr_level.loc[:, (df_3alloc_mgr_level != 0).any()].T
        df_3alloc_mgr_level = df_3alloc_mgr_level.rename_axis('Code')
        df_3alloc_mgr_level = df_3alloc_mgr_level.merge(
            Local_Portfolio.df_productList[['Name', 'G0', 'G1', 'G2', 'G3', 'G4', 'G5', 'G6', 'Type']], on='Code')
        df_3alloc_mgr_level = df_3alloc_mgr_level.rename(columns={dt_end_date: 'Current Weight'})

        df_3alloc_holding_level = df_3alloc_mgr_level
        underlying_df = []
        # Find if any of the held investments - are also available in the dataset as products with look through holdings
        for index, value in enumerate(df_3alloc_holding_level.index):
            if value in availablePortfolios:
                Underlying_Portfolio = All_Portfolios[availablePortfolios.index(value)]

                underlying_df = Underlying_Portfolio.df_L3_w.loc[dt_end_date:dt_end_date,
                                Underlying_Portfolio.df_L3_w.columns.isin(Product_List)].tail(1)
                underlying_df = underlying_df.loc[:, (underlying_df != 0).any()].T
                underlying_df = underlying_df.rename_axis('Code')

                underlying_df = underlying_df.merge(
                    Local_Portfolio.df_productList[['Name', 'G0', 'G1', 'G2', 'G3', 'G4', 'G5', 'G6', 'Type']], on='Code')
                underlying_df = underlying_df.rename(columns={dt_end_date: 'Current Weight'})

                # Find and print the 'Current Weight' in df_3alloc_holding_level
                parent_weight_value = df_3alloc_holding_level.loc[value, 'Current Weight']

                # Multiply each value in 'Current Weight' column of underlying_df
                underlying_df['Current Weight'] *= (parent_weight_value / 100)

                # Remove the matched row from df_3alloc_holding_level
                df_3alloc_holding_level = df_3alloc_holding_level.drop(index=value)

                # Merge all rows from underlying_df into df_3alloc_holding_level
                df_3alloc_holding_level = pd.concat([df_3alloc_holding_level, underlying_df])


        return df_3alloc_sleeves, df_3alloc_BMsleeves, df_3alloc_sleeve_ranges, df_3alloc_OWUW, df_3alloc_weights, df_3alloc_mgr_level, df_3alloc_holding_level

    except Exception as e:
        print(f"An error occurred in f_create_FILL_3 {e}")
        # Handle the error as needed
        return []  # or any default return value


def f_FILL_3Aequity(Local_Portfolio, BM_SharesUniverse):

    try:

        BM_SharesUniverse_latest = BM_SharesUniverse.df_L3_w.loc[dt_end_date:dt_end_date,
                                   BM_SharesUniverse.df_L3_w.columns.isin(Product_List)].tail(1)
        BM_SharesUniverse_latest = BM_SharesUniverse_latest.loc[:, (BM_SharesUniverse_latest != 0).any()].T
        BM_SharesUniverse_latest = BM_SharesUniverse_latest.rename_axis('Code')
        BM_SharesUniverse_latest = BM_SharesUniverse_latest.merge(BM_SharesUniverse.df_productList[
                                                                      ['Name', 'G0', 'G1', 'G2', 'G3', 'G4', 'G5', 'G6', 'Type',
                                                                       'LastPrice', 'MarketCap', 'BasicEPS',
                                                                       'DividendperShare-Net', 'TotalAssets',
                                                                       'TotalLiabilities', 'GrowthofNetIncome(%)',
                                                                       'GrowthofNetSales(%)',
                                                                       'GrowthofFreeCashFlow(%)', 'ReturnonTotalEquity(%)',
                                                                       'PayoutRatio', 'TotalDebt/TotalAssets',
                                                                       'NetProfitMargin(%)', 'InterestCover(EBIT)',
                                                                       'ShortSell%', 'PE_Ratio', 'EarningsYield',
                                                                       'PriceBook']],
                                                                  on='Code')
        BM_SharesUniverse_latest = BM_SharesUniverse_latest.rename(columns={dt_end_date: 'Current Weight'})

        df_3Aequity_detail_0 = Local_Portfolio.df_L3_w.loc[dt_end_date:dt_end_date,
                               Local_Portfolio.df_L3_w.columns.isin(Product_List)].tail(1)
        df_3Aequity_detail_0 = df_3Aequity_detail_0.loc[:, (df_3Aequity_detail_0 != 0).any()].T
        df_3Aequity_detail_0 = df_3Aequity_detail_0.rename_axis('Code')
        df_3Aequity_detail_0 = df_3Aequity_detail_0.merge(Local_Portfolio.df_productList[
                                                              ['Name', 'G0', 'G1', 'G2', 'G3', 'G4', 'G5', 'G6', 'Type',
                                                               'LastPrice', 'MarketCap', 'BasicEPS',
                                                               'DividendperShare-Net', 'TotalAssets', 'TotalLiabilities',
                                                               'GrowthofNetIncome(%)', 'GrowthofNetSales(%)',
                                                               'GrowthofFreeCashFlow(%)', 'ReturnonTotalEquity(%)',
                                                               'PayoutRatio', 'TotalDebt/TotalAssets',
                                                               'NetProfitMargin(%)', 'InterestCover(EBIT)', 'ShortSell%',
                                                               'PE_Ratio', 'EarningsYield', 'PriceBook']], on='Code')
        df_3Aequity_detail_0 = df_3Aequity_detail_0.rename(columns={dt_end_date: 'Current Weight'})

        # Find if any of the held investments - are also available in the dataset as products with look through holdings

        underlying_df_3A_1 = []
        for index, value in enumerate(df_3Aequity_detail_0.index):
            if value in availablePortfolios:
                Underlying_Portfolio = All_Portfolios[availablePortfolios.index(value)]

                underlying_df_3A_1 = Underlying_Portfolio.df_L3_w.loc[dt_end_date:dt_end_date,
                                     Underlying_Portfolio.df_L3_w.columns.isin(Product_List)].tail(1)
                underlying_df_3A_1 = underlying_df_3A_1.loc[:, (underlying_df_3A_1 != 0).any()].T
                underlying_df_3A_1 = underlying_df_3A_1.rename_axis('Code')

                underlying_df_3A_1 = underlying_df_3A_1.merge(
                    Local_Portfolio.df_productList[
                        ['Name', 'G0', 'G1', 'G2', 'G3', 'G4', 'G5', 'G6', 'Type', 'LastPrice', 'MarketCap', 'BasicEPS',
                         'DividendperShare-Net', 'TotalAssets', 'TotalLiabilities',
                         'GrowthofNetIncome(%)', 'GrowthofNetSales(%)',
                         'GrowthofFreeCashFlow(%)', 'ReturnonTotalEquity(%)', 'PayoutRatio', 'TotalDebt/TotalAssets',
                         'NetProfitMargin(%)', 'InterestCover(EBIT)', 'ShortSell%',
                         'PE_Ratio', 'EarningsYield', 'PriceBook']], on='Code')
                underlying_df_3A_1 = underlying_df_3A_1.rename(columns={dt_end_date: 'Current Weight'})

                # Find and print the 'Current Weight' in filtered_df_3_7
                parent_weight_value = df_3Aequity_detail_0.loc[value, 'Current Weight']

                # Multiply each value in 'Current Weight' column of underlying_df_3_7
                underlying_df_3A_1['Current Weight'] *= (parent_weight_value / 100)

                # Remove the matched row from filtered_df_3_7
                df_3Aequity_detail_0 = df_3Aequity_detail_0.drop(index=value)

                # Merge all rows from underlying_df_3_7 into filtered_df_3_7
                df_3Aequity_detail_0 = pd.concat([df_3Aequity_detail_0, underlying_df_3A_1])

        # filter for only Listed Securities
        filtered_df_3A_1 = df_3Aequity_detail_0[(df_3Aequity_detail_0['Type'] == 'ASXStock')].copy()
        assetClassWeight = filtered_df_3A_1['Current Weight'].sum()

        if assetClassWeight != 0:
            filtered_df_3A_1['Current Weight'] = filtered_df_3A_1['Current Weight'] / (assetClassWeight / 100)
        else:
            # Handle the case where assetClassWeight is zero
            filtered_df_3A_1['Current Weight'] = np.nan

        filtered_df_3A_2 = filtered_df_3A_1[(filtered_df_3A_1['Type'] == 'IEQStock')].copy()
        assetClassWeight = filtered_df_3A_1['Current Weight'].sum()

        if assetClassWeight != 0:
            filtered_df_3A_2['Current Weight'] = filtered_df_3A_2['Current Weight'] / (assetClassWeight / 100)
        else:
            # Handle the case where assetClassWeight is zero
            filtered_df_3A_2['Current Weight'] = np.nan

        # Portfolio and Benchmark Averages
        def f_calculate_normalized_percentile(selected_avg, bm_avg, data, metric):
            # Calculate the percentile of the selected value in the equal-weighted data distribution
            EW_perc_selected = percentileofscore(data[metric].dropna(), selected_avg) - 50
            EW_perc_bm = percentileofscore(data[metric].dropna(), bm_avg) - 50

            # Calculate the market capitalization-weighted percentile for the selected value
            MCap_perc_bm = 0
            if EW_perc_selected > EW_perc_bm:
                MCap_perc_selected = (((EW_perc_selected - EW_perc_bm) / (100 - EW_perc_bm)) * 50)
            else:
                MCap_perc_selected = 0 - (((EW_perc_bm - EW_perc_selected) / EW_perc_bm) * 50)

            return EW_perc_selected, EW_perc_bm, MCap_perc_selected, MCap_perc_bm

        def create_valid_variable_name(measure):
            # Remove special characters and spaces
            valid_name = re.sub(r'[^a-zA-Z0-9]', '', measure)
            return f"selected_avg_{valid_name}", f"BM_avg_{valid_name}"

        # Create an empty DataFrame to store the results
        columnsfordf1 = ['Measure', 'Category', 'Selected Avg', 'Selected MCap Normalized', 'Selected EW Normalized',
                         'Benchmark Avg', 'Benchmark MCap Normalized', 'Benchmark EW Normalized']
        df_portfolioAESummary = pd.DataFrame(columns=columnsfordf1)

        # Calculate averages for each measure
        averages = {}
        for measure, category in zip(measures, measures_category):
            selected_var, bm_var = create_valid_variable_name(measure)
            selected_avg = (filtered_df_3A_1[measure].astype(float) * filtered_df_3A_1['Current Weight'] / 100).sum()
            bm_avg = (BM_SharesUniverse_latest[measure].astype(float) * BM_SharesUniverse_latest[
                'Current Weight'] / 100).sum()

            averages[selected_var] = selected_avg
            averages[bm_var] = bm_avg



            # Avoid division by zero for measures with zero benchmark average
            if bm_avg == 0:
                bm_avg = 1e-10

            EW_norm_perc_selected, EW_norm_perc_bm, MCap_norm_perc_selected, MCap_norm_perc_bm = f_calculate_normalized_percentile(
                selected_avg,
                bm_avg,
                filtered_df_3A_1,
                measure)

            # Append results to the DataFrame
            df_portfolioAESummary = pd.concat([df_portfolioAESummary, pd.DataFrame({
                'Measure': [measure],
                'Category': [category],
                'Selected Avg': [selected_avg],
                'Selected MCap Normalized': [MCap_norm_perc_selected],
                'Selected EW Normalized': [EW_norm_perc_selected],
                'Benchmark Avg': [bm_avg],
                'Benchmark MCap Normalized': [MCap_norm_perc_bm],
                'Benchmark EW Normalized': [EW_norm_perc_bm]
            })], ignore_index=True)


        # The Below Groups Holding Weights Where Multiple Products have same Look Through Exposure
        filtered_df_3A_3 = filtered_df_3A_1.copy()
        filtered_df_3A_3['Code'] = filtered_df_3A_3.index  # this keeps it as a normal column not just index
        grouped_df_3A_3 = filtered_df_3A_3.groupby('Name').agg({
            'Code': 'first',  # Include 'Code' in the aggregation
            'Current Weight': 'sum',
            'G0': 'first', 'G1': 'first', 'G2': 'first', 'G3': 'first', 'G4': 'first', 'G5': 'first', 'G6': 'first',
            'Type': 'first', 'LastPrice': 'first', 'MarketCap': 'first', 'BasicEPS': 'first', 'DividendperShare-Net': 'first',
            'GrowthofNetIncome(%)': 'first', 'GrowthofNetSales(%)': 'first', 'GrowthofFreeCashFlow(%)': 'first',
            'ReturnonTotalEquity(%)': 'first', 'PayoutRatio': 'first', 'TotalDebt/TotalAssets': 'first',
            'NetProfitMargin(%)': 'first', 'InterestCover(EBIT)': 'first', 'ShortSell%': 'first',
            'PE_Ratio': 'first', 'EarningsYield': 'first', 'PriceBook': 'first'
        }).reset_index()

        # Weight Sorted
        grouped_df_3A_3_sorted = grouped_df_3A_3.sort_values(by='Current Weight', ascending=False)

        return df_portfolioAESummary, filtered_df_3A_1, grouped_df_3A_3, grouped_df_3A_3_sorted, averages

    except Exception as e:
        print(f"An error occurred in f_create_FILL_3a {e}")
        # Handle the error as needed
        return []  # or any default return value



## MAIN LAYOUT --------

content = html.Div(id="page-content", children=[])

app.layout = html.Div([
    dcc.Location(id="url"),
    left_sidebar,
    content,
    right_sidebar
])
@app.callback(
    Output("page-content", "children"),
    [Input("url", "pathname")]
)
def render_page_content(pathname):
    if pathname == "/":
        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Select Portfolio & Analysis Settings',
                    style={'textAlign':'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),
            dbc.Row([
                dbc.Col(dbc.Card([dbc.CardHeader("Select Portfolio:", className="card-header-bold"),
                                  dbc.CardBody([

                                      html.H6('Primary Portfolio:'),
                                      dcc.Dropdown(id='portfolio-dropdown',
                                                   options=[{'label': portfolio+" : "+All_Portfolios[availablePortfolios.index(portfolio)].portfolioName,
                                                             'value': portfolio} for portfolio in
                                                            availablePortfolios],
                                                   value=Selected_Code),
                                      html.Hr(),
                                      html.H6('Alternative 1:'),
                                      dbc.Row([
                                          dbc.Col(daq.BooleanSwitch(id='portfolio-dropdown-alt1-switch', on=Alt1_Switch_On),
                                                  className="mb-3", width = 1, style={'minWidth': 120}),
                                          dbc.Col(
                                              dcc.Loading(
                                                  dcc.Dropdown(
                                                      id='portfolio-dropdown-alt1',
                                                      options=[
                                                          {'label': portfolio + " : " + All_Portfolios[
                                                              availablePortfolios.index(portfolio)].portfolioName,
                                                           'value': portfolio} for portfolio in availablePortfolios
                                                      ],
                                                      value=Alt1_Code,
                                                  ),
                                                  style={'display': 'block' if Alt1_Switch_On else 'none'}
                                                  # Conditionally set display style
                                              ),
                                          ),
                                      ]),
                                      html.Hr(),
                                      html.H6('Alternative 2:'),
                                      dbc.Row([
                                          dbc.Col(daq.BooleanSwitch(id='portfolio-dropdown-alt2-switch', on=False),
                                                  className="mb-3", width = 1, style={'minWidth': 120}),
                                          dbc.Col(dcc.Dropdown(id='portfolio-dropdown-alt2',
                                                               options=[{'label': portfolio + " : " + All_Portfolios[
                                                                   availablePortfolios.index(portfolio)].portfolioName,
                                                                         'value': portfolio} for portfolio in
                                                                        availablePortfolios],
                                                               value=Alt2_Code), className="mb-3"),
                                      ]),
                                  ]
                            )], color="primary", outline=True, style={"height": "100%"}), width=5, align="stretch", className="mb-3"),

                ], justify="center", style={"display": "flex", "flex-wrap": "wrap"}, className="mb-3"),

            html.Hr(),
            dbc.Row([
                dbc.Col(
                    dbc.Card([dbc.CardHeader("Select Attribution Grouping:", className="card-header-bold"),
                              dbc.CardBody([
                                  dbc.Row([
                                      dbc.Col(
                                          dcc.RadioItems(
                                              options=[
                                                  {'label': ' G1 - Atchison Sleeve Categories', 'value': 'G1'},
                                                  {'label': ' G2 - CFS Edge Policy', 'value': 'G2'},
                                                  {'label': ' G3 - HUB24 Policy', 'value': 'G3'},
                                                  {'label': ' G4 - Sleeve Sub-Categories', 'value': 'G4'},
                                                  {'label': ' G5 - Geography', 'value': 'G5'},
                                              ],
                                              id="group_radio", value='G1', inline=False, labelStyle={'display': 'block'}
                                          ), align="start"),
                                  ], justify="evenly", align="start", className="mb-2"),
                              ])], color="primary", outline=True, style={"height": "100%"}),
                width=3, align="stretch", className="mb-3"),

                dbc.Col(dbc.Card(
                    [dbc.CardHeader("Select Analysis Timeframe:", className="card-header-bold"), dbc.CardBody([
                        dcc.DatePickerRange(display_format='DD-MMM-YYYY', start_date=load_start_date, day_size=35,
                                            end_date=load_end_date, id='date-picker', style={"font-size": "11px"})
                    ])], color="primary", outline=True, style={"height": "100%"}), width=2, align="start",
                    className="mb-2"),

            ], justify="center", style={"display": "flex", "flex-wrap": "wrap"}, className="mb-3"),

            html.Hr(),
            dbc.Row([
                dbc.Col(dbc.Card(
                    [
                        html.Hr(),
                        html.Div(id='display-portfolio-name',
                                 style={"color": "#1DC8F2", "margin-left": "5rem"},
                                 className="sidebar-subheader"),
                        html.Div(id='display-portfolio-type',
                                 style={"color": "#1DC8F2", "margin-left": "5rem"},
                                 className="sidebar-subheader"),
                        html.Hr(),
                    ], color="primary", outline=True, style={"height": "100%"}), width=5, align="centre",
                    className="mb-2"),
            ], justify="center", style={"display": "flex", "flex-wrap": "wrap"}, className="mb-3"),


        ]
    elif pathname == "/0-Summary":
        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Summary Dashboard',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col(width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([

                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3")
        ]
    elif pathname == "/1-Performance":

        ## Populate dataframes for Page 1-Performance
        df_1perf_daily, df_1perf_total, df_1perf_tSet, df_1perf_tMESet, df_1perf_tQESet, df_1perf_rMESet = f_FILL_1perf(Selected_Portfolio)

        df_1perf_backtestSet = df_1perf_tMESet
        df_1perf_backtestSet.columns = [Selected_Code, 'SAA Benchmark', 'Peer Group', 'Inflation']
        df_1perf_backtestSet = df_1perf_backtestSet[[Selected_Code, 'Peer Group', 'Inflation']]

        ## Populate Charts for Page 1-Performance
        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Performance Benchmarking',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([
                    # Tab 1 - Performance
                    dbc.Row([
                        dbc.Tabs([
                            dbc.Tab([
                                dbc.Card([
                                    dbc.CardHeader(
                                        "Chart 1: Total Portfolio Performance - as at Month End" +
                                        Selected_Portfolio.tME_dates.loc[0, 'Date'].strftime("(%d %b %Y)")),
                                    dbc.CardBody([dcc.Graph(figure=f_create_BAR_figure(df_1perf_tMESet, 'group', None, "Return (%, %p.a.)", "Date", 450)),
                                                  html.Hr(),
                                                  dbc.Table.from_dataframe(df_1perf_tMESet.T.round(2), index=True, striped=True, bordered=True, hover=True)
                                                  ]),
                                ], color="primary", outline=True)], label="Month End Date",
                                active_label_style={"background-color": "#1DC8F2"},
                                label_style={"background-color": "#E7EAEB", "color": "#3D555E"}),
                            dbc.Tab([
                                dbc.Card([
                                    dbc.CardHeader(
                                        "Chart 1: Total Portfolio Performance - as at Quarter End" +
                                        Selected_Portfolio.tQE_dates.loc[0, 'Date'].strftime("(%d %b %Y)")),
                                    dbc.CardBody([dcc.Graph(figure=f_create_BAR_figure(df_1perf_tQESet, 'group', None, "Return (%, %p.a.)", "Date", 450)),
                                                  html.Hr(),
                                                  dbc.Table.from_dataframe(df_1perf_tQESet.T.round(2), index=True,
                                                                           striped=True, bordered=True, hover=True)
                                                  ]),
                                ], color="primary", outline=True)], label="Quarter End Date",
                                active_label_style={"background-color": "#1DC8F2"},
                                label_style={"background-color": "#E7EAEB", "color": "#3D555E"}),
                            dbc.Tab([
                                dbc.Card([
                                    dbc.CardHeader(
                                        "Chart 1: Total Portfolio Performance - as at Last Price " +
                                        Selected_Portfolio.t_dates.loc[0, 'Date'].strftime("(%d %b %Y)")),
                                    dbc.CardBody([dcc.Graph(figure=f_create_BAR_figure(df_1perf_tSet, 'group', None, "Return (%, %p.a.)", "Date", 450)),
                                                  html.Hr(),
                                                  dbc.Table.from_dataframe(df_1perf_tSet.T.round(2), index=True,
                                                                           striped=True, bordered=True, hover=True)
                                                  ]),
                                ], color="primary", outline=True)], label="To Latest Daily",
                                active_label_style={"background-color": "#1DC8F2"},
                                label_style={"background-color": "#E7EAEB", "color": "#3D555E"}),

                            dbc.Tab([
                                dbc.Card([
                                    dbc.CardHeader(

                                        "Chart 1: Total Portfolio Performance - as at " +
                                        Selected_Portfolio.rME_dates.loc[0, 'Date'].strftime("%d %b %Y")),
                                    dbc.CardBody([dcc.Graph(
                                        figure=f_create_BAR_figure(df_1perf_rMESet[[Selected_Code, 'Peer Group', 'Inflation']], 'group', None, "Return (%, %p.a.)",

                                                                   "Date", 450)),
                                                  html.Hr(),
                                                  dbc.Table.from_dataframe(df_1perf_rMESet.T.round(2), index=True,
                                                                           striped=True, bordered=True, hover=True)
                                                  ]),

                                ], color="primary", outline=True)], label="Month End Date (Reporting)",
                                active_label_style={"background-color": "#1DC8F2"},
                                label_style={"background-color": "#E7EAEB", "color": "#3D555E"}),

                            dbc.Tab([
                                dbc.Card([
                                    dbc.CardHeader(
                                        "Chart 1: Total Portfolio Performance - as at " +
                                        Selected_Portfolio.rME_dates.loc[0, 'Date'].strftime("%d %b %Y")),
                                    dbc.CardBody([dcc.Graph(
                                        figure=f_create_BAR_figure(
                                            df_1perf_backtestSet,
                                            'group', None, "Return (%, %p.a.)",
                                            "Date", 450)),
                                        html.Hr(),
                                        dbc.Table.from_dataframe(df_1perf_backtestSet.T.round(2), index=True,
                                                                 striped=True, bordered=True, hover=True)
                                    ]),
                                ], color="primary", outline=True)], label="Backtest (Reporting)",

                                active_label_style={"background-color": "#1DC8F2"},
                                label_style={"background-color": "#E7EAEB", "color": "#3D555E"}),

                        ], className="mb-3")
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 2: Portfolio Cumulative Total Returns"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_LINE_figure(10000*(1+(df_1perf_total/100)), None, "Value of $10,000 Investment ($)", "Date",
                                                            450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 3: Example Portfolio Return Chart - Daily Asset Sleeve Returns"),
                            dbc.CardBody(dcc.Graph(figure=f_create_BAR_figure(df_1perf_daily, 'stack', None, "Daily Return (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3"),
        ]

    elif pathname == "/2-Risk":
        df_2risk_drawdown = f_CalcDrawdown(
            Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date, ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL']])
        df_2risk_drawdown.columns = [Selected_Code, 'SAA Benchmark', 'Peer Group']

        # Based on 30 day Window - Daily Data annualised (252 trading days)
        df_2risk_vol30 = f_CalcRollingDailyVol(
            Selected_Portfolio.df_L3_r.loc[(dt_start_date-timedelta(days=30)):dt_end_date,
            ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL']], 30, 252) *100
        df_2risk_vol30.columns = [Selected_Code, 'SAA Benchmark', 'Peer Group']

        # Based on 90 day Window - Daily Data annualised (252 trading days)
        df_2risk_vol90 = f_CalcRollingDailyVol(
            Selected_Portfolio.df_L3_r.loc[(dt_start_date-timedelta(days=90)):dt_end_date,
            ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL']], 90, 252) * 100
        df_2risk_vol90.columns = [Selected_Code, 'SAA Benchmark', 'Peer Group']

        # Based on 1 Year Monthly data Windows - Monthly Data annualised (12 months)
        df_2risk_vol1yr = f_CalcRollingMonthlyVol(
            Selected_Portfolio.df_L3_r.loc[(dt_start_date - timedelta(days=364)):dt_end_date,
            ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL']], 12, 12) * 100
        df_2risk_vol1yr.columns = [Selected_Code, 'SAA Benchmark', 'Peer Group']

        df_2risk_vol3yr = f_CalcRollingMonthlyVol(
            Selected_Portfolio.df_L3_r.loc[(dt_start_date - timedelta(days=(3*365-1))):dt_end_date,
            ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL']], 36, 12) * 100
        df_2risk_vol3yr.columns = [Selected_Code, 'SAA Benchmark', 'Peer Group']

        df_2risk_sharpe3yr = f_CalcRollingMonthlySharpe(
            Selected_Portfolio.df_L3_r.loc[(dt_start_date - timedelta(days=(3*365-1))):dt_end_date,
            ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL']], 36, 12, 0)
        df_2risk_sharpe3yr.columns = [Selected_Code, 'SAA Benchmark', 'Peer Group']

        df_2risk_batting3yr = f_CalcRollingMonthlyBattingAverage(
            Selected_Portfolio.df_L3_r.loc[(dt_start_date - timedelta(days=(3 * 365 - 1))):dt_end_date,
            ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL']], 36) * 100
        df_2risk_batting3yr.columns = ['SAA Benchmark', 'Peer Group']

        df_2risk_drawdown3yr = f_CalcRollingDrawdown(
            Selected_Portfolio.df_L3_r.loc[(dt_start_date - timedelta(days=(3 * 365 - 1))):dt_end_date,
            ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL']], 36) -1
        df_2risk_drawdown3yr.columns = [Selected_Code, 'SAA Benchmark', 'Peer Group']

        df_2risk_TE3yr = f_CalcRollingMonthlyTrackingError(
            Selected_Portfolio.df_L3_r.loc[(dt_start_date - timedelta(days=(3 * 365 - 1))):dt_end_date,
            ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL']], 36, 12) * 100
        df_2risk_TE3yr.columns = ['SAA Benchmark', 'Peer Group']

        # Calmar Ratio
        df_2risk_calmar3yr = f_CalcRollingMonthlyReturn(
            Selected_Portfolio.df_L3_r.loc[(dt_start_date - timedelta(days=(3 * 365 - 1))):dt_end_date,
            ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL']], 36) *100
        df_2risk_calmar3yr.columns = [Selected_Code, 'SAA Benchmark', 'Peer Group']

        # Return / <Max Drawdown
        df_2risk_calmar3yr = df_2risk_calmar3yr / -(df_2risk_drawdown3yr)
        df_2risk_calmar3yr = df_2risk_calmar3yr.dropna()

        # Information Ratio
        df_2risk_IR3yr = f_CalcRollingMonthlyAlpha(
            Selected_Portfolio.df_L3_r.loc[(dt_start_date - timedelta(days=(3 * 365 - 1))):dt_end_date,
            ['P_TOTAL', 'BM_G1_TOTAL', 'Peer_TOTAL']], 36) *100
        df_2risk_IR3yr.columns = ['SAA Benchmark', 'Peer Group']
        # Return / Tracking Error
        df_2risk_IR3yr = df_2risk_IR3yr / (df_2risk_TE3yr)

        #Create Summary Table of Latest Risk Measures
        last_df_2_2 = df_2risk_vol30.iloc[-1].copy()
        last_df_2_2['Risk Measure'] = '30 Day Volatility'
        last_df_2_3 = df_2risk_vol90.iloc[-1].copy()
        last_df_2_3['Risk Measure'] = '90 Day Volatility'
        last_df_2_4 = df_2risk_vol1yr.iloc[-1].copy()
        last_df_2_4['Risk Measure'] = '12 Month Volatility'
        last_df_2_5 = df_2risk_vol3yr.iloc[-1].copy()
        last_df_2_5['Risk Measure'] = '3 Year Volatility'
        last_df_2_8 = df_2risk_drawdown3yr.iloc[-1].copy()
        last_df_2_8['Risk Measure'] = '3 Year Max Drawdown'
        last_df_2_9 = df_2risk_TE3yr.iloc[-1].copy()
        last_df_2_9['Risk Measure'] = '3 Year Tracking Error'
        last_df_2_6 = df_2risk_sharpe3yr.iloc[-1].copy()
        last_df_2_6['Risk Measure'] = '3 Year Sharpe Ratio'
        last_df_2_10 = df_2risk_calmar3yr.iloc[-1].copy()
        last_df_2_10['Risk Measure'] = '3 Year Calmar Ratio'
        last_df_2_11 = df_2risk_IR3yr.iloc[-1].copy()
        last_df_2_11['Risk Measure'] = '3 Year Information Ratio'
        last_df_2_7 = df_2risk_batting3yr.iloc[-1].copy()
        last_df_2_7['Risk Measure'] = '3 Year Batting Average'

        df_2risk_summary = pd.concat([last_df_2_2, last_df_2_3, last_df_2_4, last_df_2_5,
                                     last_df_2_8, last_df_2_6, last_df_2_10, last_df_2_11,
                                     last_df_2_9, last_df_2_7], axis=1).T

        df_2risk_summary = df_2risk_summary.set_index('Risk Measure')
        df_2risk_summary = df_2risk_summary.apply(pd.to_numeric, errors='coerce')
        df_2risk_summary = df_2risk_summary.round(2)

        ## Populate Charts for Page 2-Risk
        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Risk Analysis',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([

                    # Tab 2 - Risk
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader(
                                "Table 1: Portfolio Risk Metrics"),
                            dbc.CardBody([dbc.Table.from_dataframe(df_2risk_summary, index=True,
                                                                   striped=True, bordered=True, hover=True,
                                                                   )
                                          ]),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 1: Portfolio 30 Day Rolling Volatility (%p.a.)"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_2risk_vol30, None, "Rolling Vol (%p.a.)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 2: Portfolio 90 Day Rolling Volatility (%p.a.)"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_2risk_vol90, None, "Rolling Vol (%p.a.)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 3: Portfolio 12 Month Rolling Volatility (%p.a.)"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_2risk_vol1yr, None, "Rolling Vol (%p.a.)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 4: Portfolio 36 Month Rolling Volatility (%p.a.)"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_2risk_vol3yr, None, "Rolling Vol (%p.a.)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 5: Portfolio Drawdown Analysis"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_2risk_drawdown, None, "Drawdown Return (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 6: Portfolio 3 Year Rolling Max Drawdown"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_2risk_drawdown3yr, None, "Max Rolling Drawdown (%)", "Date",450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 7: Portfolio 3 Year Rolling Batting Average"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_2risk_batting3yr, None, "Batting Average (%)", "Date",450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 8: Portfolio 3 Year Rolling Sharpe Ratio"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_2risk_sharpe3yr, None, "Sharpe Ratio", "Date",450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 9: Portfolio 3 Year Rolling Tracking Error"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_2risk_TE3yr, None, "Tracking Error (%)", "Date",450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 10: Portfolio 3 Year Rolling Calmar Ratio"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_2risk_calmar3yr, None, "Calmar Ratio", "Date",450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 11: Portfolio 3 Year Rolling Information Ratio"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_2risk_IR3yr, None, "Information Ratio", "Date",450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),
                        # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3")


        ]


    elif pathname == "/3-Allocation":

        ## Populate dataframes for Page 3-Allocation
        df_3alloc_sleeves, df_3alloc_BMsleeves, df_3alloc_sleeve_ranges, df_3alloc_OWUW, df_3alloc_weights, df_3alloc_mgr_level, df_3alloc_holding_level = f_FILL_3alloc(Selected_Portfolio)

        ## Populate Charts for Page 3-Allocation
        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Allocation / Exposure Analysis',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([

                    # Tab 3- Allocations
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 1: Current "+groupName+" Policy Ranges"),
                            dbc.CardBody(dcc.Graph(figure=f_create_RANGE_figure(df_3alloc_sleeve_ranges, "", "Weight (%)", "Asset Class", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 2: Current Allocation - Manager Level"),
                            dbc.CardBody([
                                html.Div([
                                    dcc.Graph(figure=f_create_PIE_figure(df_3alloc_sleeves, 'Current', 'GroupValue', None,450),
                                        style={'flex': '1'}),
                                    dcc.Graph(figure=f_create_PIE_figure(df_3alloc_BMsleeves, 'Current', 'GroupValue', None,450),
                                        style={'flex': '1'}),
                                ], style={'display': 'flex'}),
                            ]),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 3: Current Asset Allocation New"),
                            dbc.CardBody(dcc.Graph(figure=f_create_SUNBURST_figure(df_3alloc_mgr_level, ['G0', 'G1', 'G4', 'Name'], 'Name', 'Current Weight', 'Current Asset Allocation - Manager Level', 800))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 4: Current Asset Allocation"),
                            dbc.CardBody(dcc.Graph(figure=f_create_SUNBURST_figure(df_3alloc_mgr_level, ['G1', 'Name'], 'Name', 'Current Weight', 'Current Asset Allocation - Manager Level', 800))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 5: Current Asset Allocation - Drill Through"),
                            dbc.CardBody(dcc.Graph(figure=f_create_SUNBURST_figure(df_3alloc_holding_level, ['G1', 'Name'], 'Name', 'Current Weight', 'Current Asset Allocation - Holding Level', 800))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 6: Current Asset Allocation - Drill Through Sector"),
                            dbc.CardBody(dcc.Graph(figure=f_create_SUNBURST_figure(df_3alloc_holding_level, ['G1', 'G4', 'Name'], 'Name', 'Current Weight', 'Current Asset Allocation - Holding Level', 800))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 7: Current Asset Allocation - Drill Through Country"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_SUNBURST_figure(df_3alloc_holding_level, ['G1', 'G4', 'Name'], 'Name',
                                                                'Current Weight',
                                                                'Current Asset Allocation - Manager Level', 800))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader(
                                "Chart 8: Portfolio Sleeve Overweights/Underweights Through Time"),
                            dbc.CardBody(dcc.Graph(figure=f_create_BAR_figure(df_3alloc_OWUW, 'relative', None, "Overweight / Underweight (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                        dbc.Col(dbc.Card([
                            dbc.CardHeader(
                                "Chart 9: Asset Allocation Through Time"),
                            dbc.CardBody(dcc.Graph(figure=f_create_BAR_figure(df_3alloc_weights, 'stack', "Portfolio Sleeve Weights Through Time", "Weight (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        # Left Gutter
                        dbc.Col("", width=2, align="center", className="mb-3"),
                        # Centre Work Area
                        dbc.Col([
                            dbc.Table.from_dataframe(df_3alloc_holding_level, striped=True, bordered=True, hover=True)
                            # End of Centre Work Area
                        ], width=12, align="center", className="mb-3"),

                        # Right Gutter
                        dbc.Col("", width=2, align="center", className="mb-3"),

                    ], align="center", className="mb-3"),



                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3")

        ]


    elif pathname == "/3A-Equity":

        BM_AustShares_Portfolio = All_Portfolios[availablePortfolios.index("IOZ-AU")]
        #BM_IntShares_Portfolio = All_Portfolios[availablePortfolios.index("VGS-AU")]
        # toggle to include Aus and/or Int eq
        BM_SharesUniverse = BM_AustShares_Portfolio
        ## Populate dataframes for Page 3A-Equity
        # f_FILL_3Aequity(Selected_Portfolio, BM_SharesUniverse)

        df_3Aequity_AESummary, filtered_df_3A_1, grouped_df_3A_2, grouped_df_3A_2_sorted, averages = f_FILL_3Aequity(Selected_Portfolio, BM_SharesUniverse)

        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Equity Allocation - Factor Analysis',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 1 Drill Through Aus Eq Factor Characteristics"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_SCATTER_figure(grouped_df_3A_2, averages, "ReturnonTotalEquity(%)", "MarketCap", 'Current Weight', "G4",
                                                               None, None, None, 800, 1.5, x_range=(-150, 150)))),

                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 2 Drill Through Aus Eq Factor Characteristics"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_SCATTER_figure(grouped_df_3A_2, averages, "NetProfitMargin(%)",
                                                               "ReturnonTotalEquity(%)", 'Current Weight', "G4",
                                                               None, None, None, 800, 1.5,
                                                               x_range=(-250, 250), y_range=(-150, 150)))),

                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 3 Drill Through Aus Eq Factor Characteristics"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_SCATTER_figure(grouped_df_3A_2, averages, "PE_Ratio",
                                                               "ReturnonTotalEquity(%)", 'Current Weight', "G4",
                                                               None, None, None, 800, 1.5, y_range=(-150, 150)))),

                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 4 Drill Through Aus Eq Factor Characteristics"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_SCATTER_figure(grouped_df_3A_2, averages, "EarningsYield",
                                                               "GrowthofNetIncome(%)", 'Current Weight', "G4",
                                                               None, None, None, 800, 1.5,
                                                               x_range=(-0.5, 0.5), y_range=(-1000, 1000)))),

                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 10 - Portfolio Building Blocks"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_WATERFALL_figure(grouped_df_3A_2_sorted, 'Name', 'Current Weight', None, None,
                                                      None, 800, None, None))),

                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Portfolio Summary Factor Characteristics (Equal Weight Normalised)"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_COLORBAR_figure(df_3Aequity_AESummary, 'group', 'Measure', 'Selected EW Normalized', 'Category',
                                                           None, "Equal Weighted Normalized Score", "Factor Measure",
                                                           800))),

                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Portfolio Summary Factor Characteristics (Market Cap Weight Normalised)"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_COLORBAR_figure(df_3Aequity_AESummary, 'group', 'Measure',
                                                                'Selected MCap Normalized', 'Category',
                                                                None, "Equal Weighted Normalized Score",
                                                                "Factor Measure",
                                                                800))),

                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Portfolio Summary Factor Characteristics (EW Weight Normalised)"),
                            dbc.CardBody(dcc.Graph(figure=f_create_POLAR_figure(df_3Aequity_AESummary, Selected_Portfolio.portfolioCode, BM_AustShares_Portfolio.portfolioCode, "", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col([
                            dbc.Table.from_dataframe(df_3Aequity_AESummary, striped=True, bordered=True, hover=True)
                            # End of Centre Work Area
                        ], width=8, align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3")

        ]

    elif pathname == "/3B-Debt":

        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Debt Allocation Characteristics',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([

                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3")

        ]

    elif pathname == "/3C-Alternate":

        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Alternate Allocation Characteristics',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([

                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3")

        ]


    elif pathname == "/4-Attribution":
        df_4attrib_total = (((Selected_Portfolio.df_L3_1FAttrib.loc[dt_start_date:dt_end_date, ['P_TOTAL_G1 -- Allocation Effect',
                                                                                     'P_TOTAL_G1 -- Selection Effect']] + 1).cumprod() - 1) * 100)

        df_4attrib_equity = (((Selected_Portfolio.df_L3_1FAttrib.loc[dt_start_date:dt_end_date,
                         ['G1_Australian Shares-- Allocation Effect',
                          'G1_Australian Shares-- Selection Effect',
                          'G1_International Shares-- Allocation Effect',
                          'G1_International Shares-- Selection Effect']] + 1).cumprod() - 1) * 100)

        df_4attrib_alts = (((Selected_Portfolio.df_L3_1FAttrib.loc[dt_start_date:dt_end_date,
                             ['G1_Real Assets-- Allocation Effect',
                             'G1_Real Assets-- Selection Effect',
                             'G1_Alternatives-- Allocation Effect',
                             'G1_Alternatives-- Selection Effect']] + 1).cumprod() - 1) * 100)

        df_4attrib_def = (((Selected_Portfolio.df_L3_1FAttrib.loc[dt_start_date:dt_end_date,
                     ['G1_Long Duration-- Allocation Effect',
                      'G1_Long Duration-- Selection Effect',
                      'G1_Floating Rate-- Allocation Effect',
                      'G1_Floating Rate-- Selection Effect',
                      'G1_Cash-- Allocation Effect',
                      'G1_Cash-- Selection Effect',
                      ]] + 1).cumprod() - 1) * 100)

        df_4attrib_summary = (f_CalcReturnTable(Selected_Portfolio.df_L3_1FAttrib.loc[dt_start_date:dt_end_date,
                                            ['P_TOTAL_G1 -- Allocation Effect',
                                             'P_TOTAL_G1 -- Selection Effect',
                                             'G1_Australian Shares-- Allocation Effect',
                                             'G1_Australian Shares-- Selection Effect',
                                             'G1_International Shares-- Allocation Effect',
                                             'G1_International Shares-- Selection Effect',
                                             'G1_Real Assets-- Allocation Effect',
                                             'G1_Real Assets-- Selection Effect',
                                             'G1_Alternatives-- Allocation Effect',
                                             'G1_Alternatives-- Selection Effect',
                                             'G1_Long Duration-- Allocation Effect',
                                             'G1_Long Duration-- Selection Effect',
                                             'G1_Floating Rate-- Allocation Effect',
                                             'G1_Floating Rate-- Selection Effect',
                                             'G1_Cash-- Allocation Effect',
                                             'G1_Cash-- Selection Effect'
                                             ]],
            Selected_Portfolio.tME_dates) * 100).T

        ## Populate Charts for Page 4 Attribution
        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Multi-Period Brinson-Fachler Attribution Analysis',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([
                    # Tab 4- Attribution Analysis
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader(
                                "Chart 1: Portfolio Attribution Analysis vs Reference Portfolio"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_4attrib_total, None, "Value-Add Return (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 2: L3 SAA to TAA Attribution Analysis (Equities)"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_4attrib_equity, None, "Value-Add Return (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 3: L3 SAA to TAA Attribution Analysis (Alternatives)"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_4attrib_alts, None, "Value-Add Return (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 4: L3 SAA to TAA Attribution Analysis (Defensives)"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_4attrib_def, None, "Value-Add Return (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader(
                                "Chart 5: Portfolio Risk Metrics Chart - Daily Asset Sleeve Returns"),
                            dbc.CardBody([dbc.Table.from_dataframe(df_4attrib_summary.T.round(2), index=True,
                                                                   striped=True, bordered=True, hover=True)
                                          ]),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3")
        ]

    elif pathname == "/5-Contribution":

        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date]
        if checkData[0] > 0:
            df_5cont_sleeves_hasData = True
            df_5cont_sleeves = pd.concat([(((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date, ['P_' + groupName + '_' + n]] + 1).cumprod() - 1) * 100)
                                         for n in groupList], axis=1)
            df_5cont_sleeves.columns = groupList

        else:
            df_5cont_sleeves_hasData = False
            df_5cont_sleeves = pd.DataFrame()
            df_5cont_sleeves['Date'] = []
            df_5cont_sleeves.set_index('Date', inplace=True)
            df_5cont_sleeves['Value'] = []



        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date, ['P_' + groupName + '_' + "Australian Shares"]]
        if checkData[0] > 0:
            df_5cont_auseq_hasData = True
            df_5cont_auseq = (((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date, f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Australian Shares")] + 1).cumprod() - 1) * 100)
        else:
            df_5cont_auseq_hasData = False
            df_5cont_auseq = pd.DataFrame()
            df_5cont_auseq['Date'] = []
            df_5cont_auseq.set_index('Date', inplace=True)
            df_5cont_auseq['Value'] = []

        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date, ['P_' + groupName + '_' + "International Shares"]]
        if checkData[0] > 0:
            df_5cont_inteq_hasData = True
            df_5cont_inteq = (((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date, f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "International Shares")] + 1).cumprod() - 1) * 100)
        else:
            df_5cont_inteq_hasData = False
            df_5cont_inteq = pd.DataFrame()
            df_5cont_inteq['Date'] = []
            df_5cont_inteq.set_index('Date', inplace=True)
            df_5cont_inteq['Value'] = []

        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date, ['P_' + groupName + '_' + "Real Assets"]]
        if checkData[0] > 0:
            df_5cont_real_hasData = True
            df_5cont_real = (((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date, f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Real Assets")] + 1).cumprod() - 1) * 100)
        else:
            df_5cont_real_hasData = False
            df_5cont_real = pd.DataFrame()
            df_5cont_real['Date'] = []
            df_5cont_real.set_index('Date', inplace=True)
            df_5cont_real['Value'] = []

        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date, ['P_' + groupName + '_' + "Alternatives"]]
        if checkData[0] > 0:
            df_5cont_alts_hasData = True
            df_5cont_alts = (((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date, f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Alternatives")] + 1).cumprod() - 1) * 100)
        else:
            df_5cont_alts_hasData = False
            df_5cont_alts = pd.DataFrame()
            df_5cont_alts['Date'] = []
            df_5cont_alts.set_index('Date', inplace=True)
            df_5cont_alts['Value'] = []

        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date, ['P_' + groupName + '_' + "Long Duration"]]
        if checkData[0] > 0:
            df_5cont_duration_hasData = True
            df_5cont_duration = (((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date, f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Long Duration")] + 1).cumprod() - 1) * 100)
        else:
            df_5cont_duration_hasData = False
            df_5cont_duration = pd.DataFrame()
            df_5cont_duration['Date'] = []
            df_5cont_duration.set_index('Date', inplace=True)
            df_5cont_duration['Value'] = []

        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date, ['P_' + groupName + '_' + "Floating Rate"]]
        if checkData[0] > 0:
            df_5cont_floating_hasData = True
            df_5cont_floating = (((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date, f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Floating Rate")] + 1).cumprod() - 1) * 100)
        else:
            df_5cont_floating_hasData = False
            df_5cont_floating = pd.DataFrame()
            df_5cont_floating['Date'] = []
            df_5cont_floating.set_index('Date', inplace=True)
            df_5cont_floating['Value'] = []

        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date, ['P_' + groupName + '_' + "Cash"]]
        if checkData[0] > 0:
            df_5cont_cash_hasData = True
            df_5cont_cash = (((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date, f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Cash")] + 1).cumprod() - 1) * 100)
        else:
            df_5cont_cash_hasData = False
            df_5cont_cash = pd.DataFrame()
            df_5cont_cash['Date'] = []
            df_5cont_cash.set_index('Date', inplace=True)
            df_5cont_cash['Value'] = []

        ## Populate Charts for Page 5 Contribution
        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Multi-Period Weighted Return Contribution Analysis',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([
                    # Tab 6- Underlying Detail Analysis
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 1: Asset Sleeve Weighted Return Contributions"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_LINE_figure(df_5cont_sleeves, None, "Cumulative Return (%)", "Date",
                                                            450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                        # Conditionally include the chart for Australian Shares Sleeve if data is available
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 2: Australian Shares Sleeve - Contributions"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_LINE_figure(df_5cont_auseq, None, "Cumulative Return (%)", "Date",
                                                            450))),
                        ], color="primary", outline=True), align="center",
                            className="mb-3") if df_5cont_auseq_hasData else None,
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 3: International Shares Sleeve - Contributions"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_LINE_figure(df_5cont_inteq, None, "Cumulative Return (%)", "Date",
                                                            450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 4: Real Assets Sleeve - Contributions"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_LINE_figure(df_5cont_real, None, "Cumulative Return (%)", "Date",
                                                            450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 5: Alternatives Sleeve - Contributions"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_LINE_figure(df_5cont_alts, None, "Cumulative Return (%)", "Date",
                                                            450))),
                        ], color="primary", outline=True), align="center",
                            className="mb-3") if df_5cont_alts_hasData else None,
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 6: Long Duration Sleeve - Contributions"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_LINE_figure(df_5cont_duration, None, "Cumulative Return (%)", "Date",
                                                            450))),
                        ], color="primary", outline=True), align="center",
                            className="mb-3") if df_5cont_duration_hasData else None,
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 7: Floating Rate Sleeve - Contributions"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_LINE_figure(df_5cont_floating, None, "Cumulative Return (%)", "Date",
                                                            450))),
                        ], color="primary", outline=True), align="center",
                            className="mb-3") if df_5cont_floating_hasData else None,
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 8: Cash Sleeve - Contributions"),
                            dbc.CardBody(dcc.Graph(
                                figure=f_create_LINE_figure(df_5cont_cash, None, "Cumulative Return (%)", "Date",
                                                            450))),
                        ], color="primary", outline=True), align="center",
                            className="mb-3") if df_5cont_cash_hasData else None,
                    ], align="center", className="mb-3"),
                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3")
        ]
    elif pathname == "/6-Component":

        df_6comp_sleeves = pd.concat([(((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date, ['P_' + groupName + '_' + n]] + 1).cumprod() - 1) * 100)
                                     for n in groupList], axis=1)
        df_6comp_sleeves.columns = groupList
        df_6comp_auseq = (((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date, f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Australian Shares")] + 1).cumprod() - 1) * 100)
        df_6comp_inteq = (((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date, f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "International Shares")] + 1).cumprod() - 1) * 100)
        df_6comp_real = (((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date, f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Real Assets")] + 1).cumprod() - 1) * 100)
        df_6comp_alts = (((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date, f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Alternatives")] + 1).cumprod() - 1) * 100)
        df_6comp_duration = (((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date, f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Long Duration")] + 1).cumprod() - 1) * 100)
        df_6comp_floating = (((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date, f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Floating Rate")] + 1).cumprod() - 1) * 100)
        df_6comp_cash = (((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date, f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Cash")] + 1).cumprod() - 1) * 100)

        ## Populate Charts for Page 6 Component
        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Sector Sleeve - Look Through Component Analysis',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([
                    # Tab 5- Contribution Analysis
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 1: Asset Sleeve Performance"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_6comp_sleeves, None, "Cumulative Return (%)", "Date", 600))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 2: Australian Shares Sleeve - Underlying Components"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_6comp_auseq, None, "Cumulative Return (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader(
                                "Chart 3: International Shares Sleeve - Underlying Components"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_6comp_inteq, None, "Cumulative Return (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 4: Real Assets Sleeve - Underlying Components"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_6comp_real, None, "Cumulative Return (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 5: Alternatives Sleeve - Underlying Components"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_6comp_alts, None, "Cumulative Return (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 6: Long Duration Sleeve - Underlying Components"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_6comp_duration, None, "Cumulative Return (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 7: Floating Rate Sleeve - Underlying Components"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_6comp_floating, None, "Cumulative Return (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),
                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 8: Cash - Underlying Components"),
                            dbc.CardBody(dcc.Graph(figure=f_create_LINE_figure(df_6comp_cash, None, "Cumulative Return (%)", "Date", 450))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3")
        ]


    elif pathname == "/10-ESG":
        listq_10_1 = f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Australian Shares")
        filtered_df_10_1 = Selected_Portfolio.df_productList.loc[listq_10_1, ["Name", "G1", "G2", "G3", "G4", "Type",
                                                                              "E-score", "S-score", "G-score",
                                                                              "ESG-score", "Controversy-score"]]

        ## Populate Charts for Page 10 ESG
        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Portfolio ESG / Controversy Analysis',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([
                    dbc.Table.from_dataframe(filtered_df_10_1, striped=True, bordered=True, hover=True)
                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3")

        ]
    elif pathname == "/11-Fees":

        ## Populate Charts for Page 11 Fees
        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Portfolio Fee Analysis',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([

                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3")

        ]
    elif pathname == "/20-Markets":

        ## Populate Charts for Page 20 Markets
        return [
            html.Div(style={'height': '2rem'}),
            html.H2('General Market Valuation Analysis',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([

                    dbc.Row([
                        dbc.Col(dbc.Card([
                            dbc.CardHeader("Chart 1: US Interest Rates"),
                            dbc.CardBody(dcc.Graph(figure=f_create_3DSURFACE_figure(Selected_Portfolio.df_Eco_USInterestRates, "US Interest Rates", "Interest Rate", "Term", "Date", 800))),
                        ], color="primary", outline=True), align="center", className="mb-3"),
                    ], align="center", className="mb-3"),

                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3")

        ]

    elif pathname == "/21-Reports":

        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Report Generator',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([

                    dbc.Row([
                        dbc.Col(
                        dbc.Card([dbc.CardHeader("Select Which Sections To Output To Report:", className="card-header-bold"),
                                  dbc.CardBody([
                                      dbc.Row([
                                          dbc.Col([
                                              dcc.RadioItems(
                                                  options=[
                                                      {'label': ' 1 - 2-Page Summary Portfolio Analysis', 'value': 'R1'},
                                                      {'label': ' 2 - Detailed Portfolio Analysis', 'value': 'R2'},
                                                      {'label': ' 3 - Financial Markets Analysis', 'value': 'R3'},
                                                      {'label': ' 4 - 4th Type Report', 'value': 'R4'},
                                                      {'label': ' 5 - 5th Type Report', 'value': 'R5'},
                                                  ],
                                                  id="group_radio_2101", inline=False, labelStyle={'display': 'block'},
                                                  value='R1',
                                              ),
                                              html.Hr(),
                                              dbc.Button("Generate Report", id="btn_generate_report", color="primary"),
                                              html.Div(id='display-report-name', children="",
                                                       style={"color": "#1DC8F2", "margin-left": "5rem"},
                                                       className="sidebar-subheader")],
                                              align="start"
                                          ),

                                      ], justify="evenly", align="start", className="mb-2"),
                                  ])
                            ], color="primary", outline=True, style={"height": "100%"}),
                        width=8, align="stretch", className="mb-3"),
                    ], justify="center", style={"display": "flex", "flex-wrap": "wrap"}, className="mb-3"),

                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3"),
        ]


    elif pathname == "/22-Download":   # Not linked now - but keep - as the logica automatically saves to users machine

        SAVEDIR = "./Outputs/" + Selected_Code
        CHECK_FOLDER = os.path.isdir(SAVEDIR)
        if not CHECK_FOLDER: os.makedirs(SAVEDIR)
        CHECK_FOLDER = os.path.isdir(SAVEDIR+"/Charts")
        if not CHECK_FOLDER: os.makedirs(SAVEDIR+"/Charts")

        print("**** Atchison Analytics Dash App Has Downloaded A Report to: " +SAVEDIR)

        #1 FILL
        df_1perf_daily, df_1perf_total, df_1perf_tSet, df_1perf_tMESet, df_1perf_tQESet, df_1perf_rMESet = f_FILL_1perf(Selected_Portfolio)

        df_1perf_backtestSet = df_1perf_tMESet
        df_1perf_backtestSet.columns = [Selected_Code, 'SAA Benchmark', 'Peer Group', 'Inflation']
        df_1perf_backtestSet = df_1perf_backtestSet[[Selected_Code, 'Peer Group', 'Inflation']]


        # 3 FILL
        df_3alloc_sleeves, df_3alloc_BMsleeves, df_3alloc_sleeve_ranges, df_3alloc_OWUW, df_3alloc_weights, df_3alloc_mgr_level, df_3alloc_holding_level = f_FILL_3alloc(
            Selected_Portfolio)

        #3A FILL
        BM_AustShares_Portfolio = All_Portfolios[availablePortfolios.index("IOZ-AU")]
        BM_SharesUniverse = BM_AustShares_Portfolio
        df_3Aequity_AESummary, filtered_df_3A_1, grouped_df_3A_2, grouped_df_3A_2_sorted, averages = f_FILL_3Aequity(
            Selected_Portfolio, BM_SharesUniverse)

        df_2risk_drawdown = f_CalcDrawdown(
            Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date, ['P_TOTAL', 'Peer_TOTAL']])
        df_2risk_drawdown.columns = [Selected_Code, 'Peer Group']

        #5 CONTRIBUTION
        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date]
        if checkData[0] > 0:
            df_5cont_sleeves = pd.concat([(((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date,
                                             ['P_' + groupName + '_' + n]] + 1).cumprod() - 1) * 100)
                                          for n in groupList], axis=1)
            df_5cont_sleeves.columns = groupList

        else:
            df_5cont_sleeves = pd.DataFrame()

        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date, ['P_' + groupName + '_' + "Australian Shares"]]
        if checkData[0] > 0:
            df_5cont_auseq = (((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date,
                                f_AssetClassContrib(Selected_Portfolio.df_L3_contrib,
                                                    "Australian Shares")] + 1).cumprod() - 1) * 100)
        else:
            df_5cont_auseq = pd.DataFrame()

        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date, ['P_' + groupName + '_' + "International Shares"]]
        if checkData[0] > 0:
            df_5cont_inteq = (((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date,
                                f_AssetClassContrib(Selected_Portfolio.df_L3_contrib,
                                                    "International Shares")] + 1).cumprod() - 1) * 100)
        else:
            df_5cont_inteq = pd.DataFrame()

        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date, ['P_' + groupName + '_' + "Real Assets"]]
        if checkData[0] > 0:
            df_5cont_real = (((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date,
                               f_AssetClassContrib(Selected_Portfolio.df_L3_contrib,
                                                   "Real Assets")] + 1).cumprod() - 1) * 100)
        else:
            df_5cont_real = pd.DataFrame()

        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date, ['P_' + groupName + '_' + "Alternatives"]]
        if checkData[0] > 0:
            df_5cont_alts = (((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date,
                               f_AssetClassContrib(Selected_Portfolio.df_L3_contrib,
                                                   "Alternatives")] + 1).cumprod() - 1) * 100)
        else:
            df_5cont_alts = pd.DataFrame()

        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date, ['P_' + groupName + '_' + "Long Duration"]]
        if checkData[0] > 0:
            df_5cont_duration = (((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date,
                                   f_AssetClassContrib(Selected_Portfolio.df_L3_contrib,
                                                       "Long Duration")] + 1).cumprod() - 1) * 100)
        else:
            df_5cont_duration = pd.DataFrame()

        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date, ['P_' + groupName + '_' + "Floating Rate"]]
        if checkData[0] > 0:
            df_5cont_floating = (((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date,
                                   f_AssetClassContrib(Selected_Portfolio.df_L3_contrib,
                                                       "Floating Rate")] + 1).cumprod() - 1) * 100)
        else:
            df_5cont_floating = pd.DataFrame()

        checkData = Selected_Portfolio.df_L3_w.loc[dt_end_date, ['P_' + groupName + '_' + "Cash"]]
        if checkData[0] > 0:
            df_5cont_cash = (((Selected_Portfolio.df_L3_contrib.loc[dt_start_date:dt_end_date,
                               f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Cash")] + 1).cumprod() - 1) * 100)
        else:
            df_5cont_cash = pd.DataFrame()

        df_6comp_sleeves = pd.concat([(((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date,
                                         ['P_' + groupName + '_' + n]] + 1).cumprod() - 1) * 100)
                                      for n in groupList], axis=1)
        df_6comp_sleeves.columns = groupList
        df_6comp_auseq = (((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date,
                            f_AssetClassContrib(Selected_Portfolio.df_L3_contrib,
                                                "Australian Shares")] + 1).cumprod() - 1) * 100)
        df_6comp_inteq = (((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date,
                            f_AssetClassContrib(Selected_Portfolio.df_L3_contrib,
                                                "International Shares")] + 1).cumprod() - 1) * 100)
        df_6comp_real = (((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date,
                           f_AssetClassContrib(Selected_Portfolio.df_L3_contrib,
                                               "Real Assets")] + 1).cumprod() - 1) * 100)
        df_6comp_alts = (((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date,
                           f_AssetClassContrib(Selected_Portfolio.df_L3_contrib,
                                               "Alternatives")] + 1).cumprod() - 1) * 100)
        df_6comp_duration = (((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date,
                               f_AssetClassContrib(Selected_Portfolio.df_L3_contrib,
                                                   "Long Duration")] + 1).cumprod() - 1) * 100)
        df_6comp_floating = (((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date,
                               f_AssetClassContrib(Selected_Portfolio.df_L3_contrib,
                                                   "Floating Rate")] + 1).cumprod() - 1) * 100)
        df_6comp_cash = (((Selected_Portfolio.df_L3_r.loc[dt_start_date:dt_end_date,
                           f_AssetClassContrib(Selected_Portfolio.df_L3_contrib, "Cash")] + 1).cumprod() - 1) * 100)

        ## Populate Charts for Page 21 Reports
        return [

            html.Div(style={'height': '2rem'}),
            html.H2('Automated HTML Downloader',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([

                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3"),


            f_create_BAR_figure(
                df_1perf_rMESet[[Selected_Code, 'Peer Group', 'Inflation']],
                'group', None, "Return (%, %p.a.)",
                "Date", 450).write_html(SAVEDIR + "/Charts/" + '1_Performance-Main.html'),

            f_create_BAR_figure(
                df_1perf_backtestSet,
                'group', None, "Return (%, %p.a.)",
                "Date", 450).write_html(SAVEDIR + "/Charts/" + '1_Performance-Bar-Backtest.html'),

            f_create_BAR_figure(df_1perf_rMESet[[Selected_Code, 'Peer Group', 'Inflation']], 'group', None,
                                "Return (%, %p.a.)",
                                "Date", 450).write_html(SAVEDIR + "/Charts/" + '1_Performance-Daily.html'),

            f_create_LINE_figure(10000 * (1 + (df_1perf_total[[Selected_Code, 'Peer Group', 'Objective']] / 100)), None,
                                                    "Value of $10,000 Investment ($)", "Date",
                                                    450).write_html(SAVEDIR + "/Charts/" + '1_Performance-Cum.html'),

            f_create_BAR_figure(df_1perf_daily, 'stack', None, "Daily Return (%)", "Date", 450).write_html(SAVEDIR + "/Charts/" + '1_Performance-Bar.html'),

            f_create_LINE_figure(df_2risk_drawdown, None, "Drawdown Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '2_Drawdown.html'),

            f_create_RANGE_figure(df_3alloc_sleeve_ranges, "", "Weight (%)", "Asset Class", 450).write_html(SAVEDIR + "/Charts/" + '3_Alloc_Ranges.html'),


            f_create_SUNBURST_figure(df_3alloc_mgr_level, ['G0', 'G1', 'G4', 'Name'], 'Name',
                                                    'Current Weight', '',
                                                    800).write_html(SAVEDIR + "/Charts/" + '3_Alloc_Mgr_Level_1.html'),

            f_create_SUNBURST_figure(df_3alloc_mgr_level, ['G1', 'Name'], 'Name', 'Current Weight',
                                                    '',
                                                    800).write_html(SAVEDIR + "/Charts/" + '3_Alloc_Mgr_Level_2.html'),

            f_create_SUNBURST_figure(df_3alloc_holding_level, ['G1', 'Name'], 'Name', 'Current Weight',
                                                    '',
                                                    800).write_html(SAVEDIR + "/Charts/" + '3_Alloc_Holding_Level_1.html'),

            f_create_SUNBURST_figure(df_3alloc_holding_level, ['G1', 'G4', 'Name'], 'Name',
                                                    'Current Weight', '',
                                                    800).write_html(SAVEDIR + "/Charts/" + '3_Alloc_Holding_Level_2.html'),

            f_create_BAR_figure(df_3alloc_OWUW, 'relative', None, "Overweight / Underweight (%)", "Date", 450).write_html(SAVEDIR + "/Charts/" + '3_Allocation_OW_UW.html'),

            f_create_BAR_figure(df_3alloc_weights, 'stack', "Portfolio Sleeve Weights Through Time",
                                       "Weight (%)", "Date", 450).write_html(SAVEDIR + "/Charts/" + '3_Allocation_History.html'),

            f_create_SCATTER_figure(grouped_df_3A_2, averages, "ReturnonTotalEquity(%)", "MarketCap", 'Current Weight',
                                    "G4",None, None, None, 800,
                                    1.5, x_range=(-150, 150)).write_html(SAVEDIR + "/Charts/" + '3a_Aus_Equity_Financial_Ratios_1.html'),

            f_create_SCATTER_figure(grouped_df_3A_2, averages, "NetProfitMargin(%)",
                                    "ReturnonTotalEquity(%)", 'Current Weight', "G4",
                                    None, None, None, 800, 1.5, x_range=(-250, 250), y_range=(-150, 150)).write_html(SAVEDIR + "/Charts/" + '3a_Aus_Equity_Financial_Ratios_2.html'),

            f_create_SCATTER_figure(grouped_df_3A_2, averages, "PE_Ratio",
                                    "ReturnonTotalEquity(%)", 'Current Weight', "G4",
                                    None, None, None, 800, 1.5, y_range=(-150, 150)).write_html(SAVEDIR + "/Charts/" + '3a_Aus_Equity_Financial_Ratios_3.html'),

            f_create_SCATTER_figure(grouped_df_3A_2, averages, "EarningsYield",
                                    "GrowthofNetIncome(%)", 'Current Weight', "G4",
                                    None, None, None, 800, 1.5, x_range=(-0.50, 0.50), y_range=(-1000, 1000)).write_html(SAVEDIR + "/Charts/" + '3a_Aus_Equity_Financial_Ratios_4.html'),

            f_create_WATERFALL_figure(grouped_df_3A_2_sorted, 'Name', 'Current Weight', None, None,
                                      None, 800, None, None).write_html(SAVEDIR + "/Charts/" + '3a_Aus_Equity_Alloc_Waterfall.html'),


            f_create_3DSURFACE_figure(Selected_Portfolio.df_Eco_USInterestRates, "US Interest Rates", "Interest Rate",
                                      "Term", "Date", 800).write_html(SAVEDIR + "/Charts/" + '20_Eco_USInterestRates.html'),

            f_create_COLORBAR_figure(df_3Aequity_AESummary, 'group', 'Measure',
                                     'Selected MCap Normalized', 'Category',
                                     None, "Equal Weighted Normalized Score",
                                     "Factor Measure",
                                     800).write_html(SAVEDIR + "/Charts/" + '3a_Aus_Equity_Factor_Ratios.html'),

            f_create_COLORBAR_figure(df_3Aequity_AESummary, 'group', 'Measure',
                                     'Selected MCap Normalized', 'Category',
                                     None, "Equal Weighted Normalized Score",
                                     "Factor Measure",
                                     800).write_image(SAVEDIR + "/Charts/" + '3a_Aus_Equity_Factor_Ratios.svg'),

            f_create_LINE_figure(df_5cont_sleeves, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '5_Sleeve_Contribs.html'),
            f_create_LINE_figure(df_5cont_auseq, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '5_AusEq_Sleeve_Contribs.html'),
            f_create_LINE_figure(df_5cont_inteq, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '5_IntEq_Sleeve_Contribs.html'),
            f_create_LINE_figure(df_5cont_real, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '5_Real_Sleeve_Contribs.html'),
            f_create_LINE_figure(df_5cont_alts, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '5_Alts_Sleeve_Contribs.html'),
            f_create_LINE_figure(df_5cont_duration, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '5_Duration_Sleeve_Contribs.html'),
            f_create_LINE_figure(df_5cont_floating, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '5_Floating_Sleeve_Contribs.html'),
            f_create_LINE_figure(df_5cont_cash, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '5_Cash_Sleeve_Contribs.html'),

            f_create_LINE_figure(df_6comp_sleeves, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '6_Sleeve_Components.html'),
            f_create_LINE_figure(df_6comp_auseq, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '6_AusEq_Sleeve_Components.html'),
            f_create_LINE_figure(df_6comp_inteq, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '6_IntEq_Sleeve_Components.html'),
            f_create_LINE_figure(df_6comp_real, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '6_Real_Sleeve_Components.html'),
            f_create_LINE_figure(df_6comp_alts, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '6_Alts_Sleeve_Components.html'),
            f_create_LINE_figure(df_6comp_duration, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '6_Duration_Sleeve_Components.html'),
            f_create_LINE_figure(df_6comp_floating, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '6_Floating_Sleeve_Components.html'),
            f_create_LINE_figure(df_6comp_cash, None, "Cumulative Return (%)", "Date", 450).write_html(
                SAVEDIR + "/Charts/" + '6_Cash_Sleeve_Components.html'),

            f_create_SUMMARY_REPORT_HTML(df_marketCommentary),

        ]
    elif pathname == "/30-Help":

        ## Populate Charts for Page 30 Help
        return [
            html.Div(style={'height': '2rem'}),
            html.H2('Need Help & Model Assumptions',
                    style={'textAlign': 'center', 'color': "#3D555E"}),
            html.Hr(),
            html.Hr(style={'border-color': "#3D555E", 'width': '70%', 'margin': 'auto auto'}),
            html.Hr(),

            dbc.Row([
                # Left Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),
                # Centre Work Area
                dbc.Col([

                    # End of Centre Work Area
                ], width=8, align="center", className="mb-3"),

                # Right Gutter
                dbc.Col("", width=2, align="center", className="mb-3"),

            ], align="center", className="mb-3")

        ]
    # If the user tries to reach a different page, return a 404 message
    return dbc.Jumbotron(
        [
            html.Div(style={'height': '2rem'}),
            html.H2("404: Not found", className="text-danger"),
            html.Hr(),
            html.P(f"Whoops the pathname {pathname} was not recognised... - Blame Jake!"),
        ]
    )


#@@@ CALL BACKS @@@@@
#@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@

# Callback to set Selected_Portfolio and update the dcc.Store with Portfolio_Code
@app.callback(
    Output('display-portfolio-code', 'children'),
    Output('display-portfolio-name', 'children'),
    Output('display-portfolio-type', 'children'),
    Output('stored-portfolio-code', 'data'),
    Output('stored-alt1-switch', 'data'),
    Output('stored-alt2-switch', 'data'),
    State('stored-portfolio-code', 'data'),
    State('url', 'pathname'),
    Input('portfolio-dropdown', 'value'),
    Input('portfolio-dropdown-alt1', 'value'),
    Input('portfolio-dropdown-alt2', 'value'),
    Input('group_radio', 'value'),
    Input('date-picker', 'start_date'),  # Add start_date input
    Input('date-picker', 'end_date'),    # Add end_date input
    Input('portfolio-dropdown-alt1-switch', 'on'),
    Input('portfolio-dropdown-alt2-switch', 'on'),
)
def update_selected_portfolio(stored_value, pathname, selected_value, alt1_value, alt2_value, group_value, tx_Start_Date, tx_End_Date, alt1_on, alt2_on):
    global Selected_Portfolio, Selected_Code, Selected_Name, Selected_Type, Alt1_Portfolio, Alt1_Code, Alt1_Name, Alt1_Name, Alt1_Type
    global Alt2_Portfolio, Alt2_Code, Alt2_Name, Alt2_Type, Group_value, dt_start_date, dt_end_date, Alt1_Switch_On, Alt2_Switch_On
    global text_Start_Date, text_End_Date # Declare global variables

    if pathname == "/":
        if selected_value in availablePortfolios:
            Selected_Portfolio = All_Portfolios[availablePortfolios.index(selected_value)]
            Selected_Code = Selected_Portfolio.portfolioCode  # Update Selected_Code
            Selected_Name = Selected_Portfolio.portfolioName
            Selected_Type = Selected_Portfolio.portfolioType
            Alt1_Portfolio = All_Portfolios[availablePortfolios.index(alt1_value)]
            Alt1_Code = Alt1_Portfolio.portfolioCode
            Alt1_Name = Alt1_Portfolio.portfolioName
            Alt1_Type = Alt1_Portfolio.portfolioType
            Alt2_Portfolio = All_Portfolios[availablePortfolios.index(alt2_value)]
            Alt2_Code = Alt2_Portfolio.portfolioCode
            Alt2_Name = Alt2_Portfolio.portfolioName
            Alt2_Type = Alt2_Portfolio.portfolioType
            Group_value = group_value
            dt_start_date = pd.to_datetime(tx_Start_Date)
            dt_end_date = pd.to_datetime(tx_End_Date)
            text_Start_Date = tx_Start_Date
            text_End_Date = tx_End_Date
            Alt1_Switch_On = alt1_on
            Alt2_Switch_On = alt2_on

            return Selected_Code, Selected_Name, Selected_Type, {'key': Selected_Code}, {'key': Alt1_Switch_On}, {'key': Alt2_Switch_On}
        else:
            return None, None, None, None, None, None
    else:
        return None, None, None, None, None, None


@app.callback(
    Output("sidebar-left-id", "style"),
    Input("pin-toggle-button", "n_clicks"),
    prevent_initial_call=True,
)
def toggle_sidebar(n_clicks):
    width = "20rem" if n_clicks % 2 == 1 else "3.5rem"
    return {"width": width}



@app.callback(
    Output('btn_generate_report', 'n_clicks'),
    [Input('btn_generate_report', 'n_clicks')],
    [State('group_radio_2101', 'value')],
    prevent_initial_call=True
)
def save_report_to_terminal(n_clicks, selected_report):
    if n_clicks is not None:

        ########################################## Create & Populate Word document #######################
        print("Current Working Directory:", os.getcwd())

        # Define Content for Word Report
        if selected_report == 'R1':
            print('Report is 2-Page Summary Portfolio Analysis')
            ########################################## Create & Populate Word document #######################
            doc = Document('./assets/Template-Short.docx')
            table_number = 1
            figure_number = 1
            # Header section
            # header_section = doc.sections[0]
            # header = header_section.header
            # header_text = header.paragraphs[0]
            # header_text.text = "ATCHISON"
            ## Heading page
            title_para = doc.add_paragraph()
            title = title_para.add_run(Selected_Name)
            title.alignment = 2
            title.font.name = 'Arial'
            title.font.size = docx.shared.Pt(18)
            title.font.color.rgb = docx.shared.RGBColor(117, 193, 4)
            sub = doc.add_paragraph()
            subsub = sub.add_run('Period to ' + dt_end_date.strftime('%d %B %Y'))
            subsub.alignment = 1
            subsub.font.name = 'Arial'
            subsub.font.size = docx.shared.Pt(14)
            subsub.font.color.rgb = docx.shared.RGBColor(117, 193, 4)

            ############################################### 1 Executive Summary ########################################
            doc.add_paragraph('Performance Summary', style='BodyStyle')



            #

            # Save Document to Local System
            SAVEDIR = "./Outputs/" + Selected_Code
            CHECK_FOLDER = os.path.isdir(SAVEDIR)
            if not CHECK_FOLDER: os.makedirs(SAVEDIR)
            docsavename = SAVEDIR + "/" + Selected_Name + " Performance Summary - " + text_End_Date + ".docx"
            doc.save(docsavename)

        elif selected_report == 'R2':
            print('Detailed Portfolio Analysis')

            ########################################## Create & Populate Word document #######################
            doc = Document('./assets/Template.docx')
            table_number = 1
            figure_number = 1
            # Header section
            # header_section = doc.sections[0]
            # header = header_section.header
            # header_text = header.paragraphs[0]
            # header_text.text = "ATCHISON"
            ## Heading page
            title_para = doc.add_paragraph()
            title = title_para.add_run(Selected_Name)
            title.alignment = 2
            title.font.name = 'Arial'
            title.font.size = docx.shared.Pt(26)
            title.font.color.rgb = docx.shared.RGBColor(117, 193, 4)
            sub = doc.add_paragraph()
            subsub = sub.add_run('Review of Investment Portfolio ' + str(dt_end_date))
            subsub.alignment = 1
            subsub.font.name = 'Arial'
            subsub.font.size = docx.shared.Pt(16)
            subsub.font.color.rgb = docx.shared.RGBColor(255, 255, 255)
            doc.add_page_break()
            ## Table of content
            ToC = doc.add_paragraph()
            toc_run = ToC.add_run('Table of Contents')
            toc_run.font.name = 'Arial'
            toc_run.font.size = docx.shared.Pt(16)
            toc_run.font.color.rgb = docx.shared.RGBColor(117, 193, 4)
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')  # creates a new element
            fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
            instrText.text = 'TOC \\o "1-5" \\h \\z \\u'  # change 1-3 depending on heading levels you need
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            fldChar3 = OxmlElement('w:t')
            fldChar3.text = "Right-click to update field."
            fldChar2.append(fldChar3)
            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')
            r_element = run._r
            r_element.append(fldChar)
            r_element.append(instrText)
            r_element.append(fldChar2)
            r_element.append(fldChar4)
            p_element = paragraph._p
            doc.add_page_break()
            ############################################### 1 Executive Summary ########################################
            doc.add_heading("Executive Summary", 1)

            #

            # Save Document to Local System
            SAVEDIR = "./Outputs/" + Selected_Code
            CHECK_FOLDER = os.path.isdir(SAVEDIR)
            if not CHECK_FOLDER: os.makedirs(SAVEDIR)
            docsavename = SAVEDIR + "/" + Selected_Name + " Quant Analysis - " + text_End_Date + ".docx"
            doc.save(docsavename)

        elif selected_report == 'R3':
            print('Financial Markets Analysis')

        else:
            return None



        f_save_report(selected_report)
        print(f'Report generated for {selected_report}.')
    return None


#text_Start_Date = load_start_date
#text_End_Date = load_end_date

# Update the CSS styles for printing to ensure A4 size
# app.css.append_css({'external_url': 'https://cdnjs.cloudflare.com/ajax/libs/font-awesome/4.7.0/css/font-awesome.min.css'})



# Run the app
if __name__ == '__main__':
    #app.run_server(debug=True)
    app.run_server(host = '0.0.0.0', port = port_number, debug=True)
    #app.run_server(host='0.0.0.0', port=port_number)